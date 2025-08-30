import os
import shutil
from typing import List, Dict

import pandas as pd
from docx import Document
from docx.shared import Cm

# Chemin du template Word par défaut (relatif au dépôt)
TEMPLATE_FILE = os.path.abspath(
    os.path.join(os.path.dirname(__file__), '..', 'Template word  Contexte éco', '1 Template Contexte éco.docx')
)

# Configuration de base pour les différents zonages. Cette structure peut
# être complétée au besoin en ajoutant de nouvelles entrées.
DEFAULT_MAPPINGS: List[Dict[str, str]] = [
    {
        'table_marker': 'TABLEAU NATURA2000',
        'sheet_prefix': 'N2000',
        'image_marker': 'CARTE NATURA2000',
        'image_name': 'Contexte éco - N2000__AE.png',
    },
    {
        'table_marker': 'TABLEAU ENS',
        'sheet_prefix': 'ENS',
        'image_marker': 'CARTE ENS',
        'image_name': 'Contexte éco - ENS__AE.png',
    },
    {
        'table_marker': 'TABLEAU APPB',
        'sheet_prefix': 'APPB',
        'image_marker': 'CARTE APPB',
        'image_name': 'Contexte éco - APPB__AE.png',
    },
]

def _replace_table(paragraph, df: pd.DataFrame) -> None:
    """Insère un tableau DataFrame avant un paragraphe et supprime le marqueur."""
    rows, cols = df.shape
    table = paragraph._parent.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        header_cells[j].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].text = "" if pd.isna(value) else str(value)
    # Déplacer le tableau avant le paragraphe puis supprimer celui-ci
    paragraph._element.addprevious(table._element)
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)

def _replace_image(paragraph, image_path: str) -> None:
    """Remplace un paragraphe par une image centrée."""
    paragraph.clear()
    run = paragraph.add_run()
    try:
        run.add_picture(image_path, width=Cm(14))
    except Exception:
        # En cas d'échec (fichier manquant, etc.), on laisse le paragraphe vide
        return

def generate_word_report(
    excel_path: str,
    images_dir: str,
    output_docx: str,
    template_path: str = TEMPLATE_FILE,
    mappings: List[Dict[str, str]] = None,
) -> str:
    """Génère un rapport Word à partir des résultats d'identification et des cartes.

    :param excel_path: Chemin du fichier Excel produit par l'analyse "ID Contexte éco".
    :param images_dir: Dossier contenant les cartes exportées (PNG).
    :param output_docx: Chemin du fichier Word à créer.
    :param template_path: Chemin du modèle Word à copier avant insertion.
    :param mappings: Liste de dictionnaires décrivant les marqueurs et fichiers associés.
    :return: Chemin du document Word généré.
    """
    if mappings is None:
        mappings = DEFAULT_MAPPINGS

    os.makedirs(os.path.dirname(output_docx), exist_ok=True)
    shutil.copy(template_path, output_docx)

    doc = Document(output_docx)
    sheets = pd.read_excel(excel_path, sheet_name=None)

    for mapping in mappings:
        # Insertion du tableau
        marker = mapping.get('table_marker')
        sheet_prefix = mapping.get('sheet_prefix')
        df = None
        if sheet_prefix:
            to_concat = [df for name, df in sheets.items() if name.startswith(sheet_prefix)]
            if to_concat:
                df = pd.concat(to_concat, ignore_index=True)
        else:
            sheet_name = mapping.get('sheet_name')
            df = sheets.get(sheet_name)
        if df is not None:
            for paragraph in doc.paragraphs:
                if marker in paragraph.text:
                    _replace_table(paragraph, df)
                    break
        # Insertion de l'image
        img_marker = mapping.get('image_marker')
        img_name = mapping.get('image_name')
        if img_marker and img_name:
            img_path = os.path.join(images_dir, img_name)
            if os.path.isfile(img_path):
                for paragraph in doc.paragraphs:
                    if img_marker in paragraph.text:
                        _replace_image(paragraph, img_path)
                        break

    doc.save(output_docx)
    return output_docx
