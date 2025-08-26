import os
import re
import threading
import tempfile
import webbrowser
from io import BytesIO
from typing import List

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter import font as tkfont
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from selenium.webdriver import ActionChains
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image

from .style_helper import StyleHelper
from .config import (
    WAIT_TILES_DEFAULT, OUTPUT_DIR_RLT, LAYERS, URL,
    IMG_WIDTH, WORD_FILENAME, COMMENT_TEMPLATE
)
from .utils import ToolTip, TextRedirector, log_with_time, load_prefs, save_prefs

def dms_to_dd(text: str) -> float:
    pat = r"(\d{1,3})[°d]\s*(\d{1,2})['m]\s*([\d\.]+)[\"s]?\s*([NSEW])"
    alt = r"(\d{1,3})\s+(\d{1,2})\s+([\d\.]+)\s*([NSEW])"
    m = re.search(pat, text, re.I) or re.search(alt, text, re.I)
    if not m:
        raise ValueError(f"Format DMS invalide : {text}")
    deg, mn, sc, hemi = m.groups()
    dd = float(deg) + float(mn)/60 + float(sc)/3600
    return -dd if hemi.upper() in ("S", "W") else dd

def add_hyperlink(paragraph, url: str, text: str, italic: bool = True):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    from docx.oxml import OxmlElement, ns
    fld_simple = OxmlElement('w:hyperlink')
    fld_simple.set(ns.qn('r:id'), r_id)

    run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    if italic:
        i = OxmlElement('w:i')
        r_pr.append(i)
    u = OxmlElement('w:u')
    u.set(ns.qn('w:val'), 'single')
    r_pr.append(u)
    run.append(r_pr)
    run_text = OxmlElement('w:t')
    run_text.text = text
    run.append(run_text)
    fld_simple.append(run)
    paragraph._p.append(fld_simple)

# =========================
# UI — Styles communs
# =========================
class RemonterLeTempsTab(ttk.Frame):
    def __init__(self, parent, style_helper: StyleHelper, prefs: dict):
        super().__init__(parent, padding=12)
        self.style_helper = style_helper
        self.prefs = prefs

        self.font_title = tkfont.Font(family="Segoe UI", size=15, weight="bold")
        self.font_sub   = tkfont.Font(family="Segoe UI", size=10)
        self.font_mono  = tkfont.Font(family="Consolas", size=9)

        self.coord_var   = tk.StringVar(value=self.prefs.get("RLT_COORD", ""))   # ex: 45°09'30" N 5°43'12" E
        self.commune_var = tk.StringVar()
        self.wait_var    = tk.DoubleVar(value=float(self.prefs.get("RLT_WAIT", WAIT_TILES_DEFAULT)))
        self.out_dir_var = tk.StringVar(value=self.prefs.get("RLT_OUT", OUTPUT_DIR_RLT))
        self.headless_var= tk.BooleanVar(value=bool(self.prefs.get("RLT_HEADLESS", False)))

        self._build_ui()

    def _build_ui(self):
        header = ttk.Frame(self, style="Header.TFrame", padding=(14, 12))
        header.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(header, text="IGN « Remonter le temps » — Capture + Word", style="Card.TLabel", font=self.font_title)\
            .grid(row=0, column=0, sticky="w")
        ttk.Label(header, text="Entrer coordonnées DMS. Générer 2×2 + commentaire.", style="Subtle.TLabel", font=self.font_sub)\
            .grid(row=1, column=0, sticky="w", pady=(4,0))
        header.columnconfigure(0, weight=1)

        # Carte paramètres
        card = ttk.Frame(self, style="Card.TFrame", padding=12)
        card.pack(fill=tk.X)
        r = 0
        ttk.Label(card, text="Coordonnées DMS (lat puis lon)", style="Card.TLabel").grid(row=r, column=0, sticky="w")
        ttk.Entry(card, textvariable=self.coord_var).grid(row=r, column=1, sticky="ew", padx=8)
        ToolTip(card, "Exemple : 45°09'30\" N 5°43'12\" E")
        r += 1

        ttk.Label(card, text="Commune détectée", style="Card.TLabel").grid(row=r, column=0, sticky="w")
        self.commune_label = ttk.Label(card, textvariable=self.commune_var, style="Card.TLabel")
        self.commune_label.grid(row=r, column=1, sticky="w", padx=8)
        r += 1

        ttk.Label(card, text="Dossier de sortie", style="Card.TLabel").grid(row=r, column=0, sticky="w")
        out_row = ttk.Frame(card, style="Card.TFrame")
        out_row.grid(row=r, column=1, sticky="ew", padx=0)
        out_row.columnconfigure(0, weight=1)
        ttk.Entry(out_row, textvariable=self.out_dir_var).grid(row=0, column=0, sticky="ew")
        ttk.Button(out_row, text="Parcourir…", command=self._pick_out_dir).grid(row=0, column=1, padx=(6,0))
        r += 1

        ttk.Label(card, text="Attente tuiles (s)", style="Card.TLabel").grid(row=r, column=0, sticky="w")
        ttk.Spinbox(card, from_=0.5, to=5.0, increment=0.1, textvariable=self.wait_var, width=6, justify="right").grid(row=r, column=1, sticky="w", padx=8)
        r += 1

        ttk.Checkbutton(card, text="Mode headless (Chrome sans fenêtre)", variable=self.headless_var, style="Card.TCheckbutton")\
            .grid(row=r, column=1, sticky="w", padx=8, pady=(4,0))
        r += 1

        card.columnconfigure(1, weight=1)

        # Actions
        act = ttk.Frame(self, style="Card.TFrame", padding=12)
        act.pack(fill=tk.X, pady=(10,0))
        self.run_btn = ttk.Button(act, text="▶ Capturer + Générer le Word", style="Accent.TButton", command=self._start_thread)
        self.run_btn.grid(row=0, column=0, sticky="w")
        obtn = ttk.Button(act, text="📂 Ouvrir le dossier de sortie", command=self._open_out_dir)
        obtn.grid(row=0, column=1, padx=(10,0)); ToolTip(obtn, "Ouvrir le dossier cible")
        gbtn = ttk.Button(act, text="🌍 Ouvrir Google Maps", command=self._open_gmaps)
        gbtn.grid(row=0, column=2, padx=(10,0)); ToolTip(gbtn, "Ouvrir Google Maps")
        self.bassin_btn = ttk.Button(act, text="💧 Bassin versant", command=self._start_bassin_thread)
        self.bassin_btn.grid(row=0, column=3, padx=(10,0)); ToolTip(self.bassin_btn, "Télécharger le bassin versant")

        # Logs
        bottom = ttk.Frame(self, style="Card.TFrame", padding=12)
        bottom.pack(fill=tk.BOTH, expand=True, pady=(10,0))
        self.status_label = ttk.Label(bottom, text="Prêt.", style="Status.TLabel")
        self.status_label.grid(row=0, column=0, sticky="w")
        bottom.columnconfigure(0, weight=1)

        log_frame = ttk.Frame(bottom, style="Card.TFrame")
        log_frame.grid(row=1, column=0, sticky="nsew", pady=(8,0))
        bottom.rowconfigure(1, weight=1)

        self.log_text = tk.Text(log_frame, height=12, wrap=tk.WORD, state='disabled',
                                bg=self.style_helper.style.lookup("Card.TFrame", "background"),
                                fg=self.style_helper.style.lookup("TLabel", "foreground"))
        self.log_text.configure(font=self.font_mono, relief="flat")
        log_scroll = ttk.Scrollbar(log_frame, orient="vertical", command=self.log_text.yview)
        self.log_text['yscrollcommand'] = log_scroll.set
        log_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.stdout_redirect = TextRedirector(self.log_text)  # dédié à l’onglet
        # Utilisation: print(..., file=self.stdout_redirect)

    # --- Actions onglet 2 ---
    def _pick_out_dir(self):
        base = self.out_dir_var.get() or OUTPUT_DIR_RLT
        d = filedialog.askdirectory(title="Choisir le dossier de sortie", initialdir=base if os.path.isdir(base) else os.path.expanduser("~"))
        if d:
            self.out_dir_var.set(d)

    def _open_out_dir(self):
        try:
            os.makedirs(self.out_dir_var.get() or OUTPUT_DIR_RLT, exist_ok=True)
            os.startfile(self.out_dir_var.get() or OUTPUT_DIR_RLT)
        except Exception as e:
            messagebox.showerror("Erreur", f"Impossible d’ouvrir le dossier : {e}")

    def _open_gmaps(self):
        url = ("https://www.google.com/maps/@45.1884514,5.711743,9768m/data=!3m1!1e3?hl=fr&entry=ttu&g_ep=EgoyMDI1MDgxOS4w"
               "IKXMDSoASAFQAw%3D%3D")
        webbrowser.open(url)

    def _start_bassin_thread(self):
        if not self.coord_var.get().strip():
            messagebox.showerror("Erreur", "Renseigner les coordonnées en DMS."); return
        self.bassin_btn.config(state="disabled")
        t = threading.Thread(target=self._run_bassin)
        t.daemon = True
        t.start()

    def _run_bassin(self):
        try:
            user_address = self.coord_var.get().strip()
            download_dir = OUT_IMG
            target_folder_name = "Bassin versant"
            target_path = os.path.join(download_dir, target_folder_name)
            os.makedirs(download_dir, exist_ok=True)

            parts = re.split(r"\s{2,}|,|\t", user_address)
            if len(parts) < 2:
                parts = re.split(r"\s+", user_address, maxsplit=1)
            if len(parts) >= 2:
                try:
                    lat_dd = dms_to_dd(parts[0]); lon_dd = dms_to_dd(parts[1])
                    print(f"[BV] Coordonnées : {lat_dd:.6f}, {lon_dd:.6f}", file=self.stdout_redirect)
                except Exception:
                    print("[BV] Format DMS invalide.", file=self.stdout_redirect)

            options = webdriver.ChromeOptions()
            options.add_argument("--log-level=3")
            options.add_experimental_option('excludeSwitches', ['enable-logging'])
            options.add_argument("--disable-extensions")
            prefs = {
                "download.default_directory": download_dir,
                "download.prompt_for_download": False,
                "profile.default_content_settings.popups": 0,
                "download.directory_upgrade": True
            }
            options.add_experimental_option("prefs", prefs)

            print("[BV] Initialisation du navigateur...", file=self.stdout_redirect)
            driver = webdriver.Chrome(options=options)
            try:
                driver.maximize_window()
            except Exception:
                pass

            url = "https://mghydro.com/watersheds/"
            print(f"[BV] Navigation vers {url}...", file=self.stdout_redirect)
            driver.get(url)
            wait = WebDriverWait(driver, 2)

            opts_button = wait.until(EC.element_to_be_clickable((By.ID, "opts_click")))
            opts_button.click(); time.sleep(1)
            downloadable_checkbox = wait.until(EC.element_to_be_clickable((By.ID, "downloadable")))
            if not downloadable_checkbox.is_selected():
                downloadable_checkbox.click()

            search_icon = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, ".leaflet-control-search .search-button")))
            search_icon.click()
            search_input = wait.until(EC.visibility_of_element_located((By.ID, "searchtext84")))
            search_input.clear(); time.sleep(0.3)
            search_input.send_keys(user_address); time.sleep(1.5)
            search_input.send_keys(Keys.ARROW_DOWN); time.sleep(0.3)
            search_input.send_keys(Keys.ENTER); time.sleep(1.5)

            map_element = wait.until(EC.presence_of_element_located((By.ID, "map")))
            ActionChains(driver).move_to_element(map_element).click().perform(); time.sleep(0.8)
            wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "div.leaflet-popup")))
            delineate_button = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, ".leaflet-popup .gobutton")))
            delineate_button.click(); time.sleep(1.5)

            downloads_button = wait.until(EC.element_to_be_clickable((By.XPATH, "//span[contains(@class, 'ui-selectmenu-text') and contains(text(), 'Watershed Boundary')]")))
            downloads_button.click(); time.sleep(0.8)
            ActionChains(driver).send_keys(Keys.ARROW_DOWN).pause(0.3).send_keys(Keys.ARROW_DOWN).pause(0.3).send_keys(Keys.ENTER).perform()
            time.sleep(1.5)

        except Exception as e:
            print(f"[BV] Erreur Selenium : {e}", file=self.stdout_redirect)
        finally:
            try:
                driver.quit()
            except Exception:
                pass

        try:
            print("[BV] Attente du téléchargement du ZIP...", file=self.stdout_redirect)
            zip_file_path = None
            wait_time = 30
            start_time = time.time()
            while time.time() - start_time < wait_time:
                zip_candidates = [f for f in os.listdir(download_dir) if f.lower().endswith('.zip') and not f.lower().endswith('.crdownload')]
                if zip_candidates:
                    zip_candidates_full = [os.path.join(download_dir, z) for z in zip_candidates]
                    zip_file_path = max(zip_candidates_full, key=os.path.getmtime)
                    size1 = os.path.getsize(zip_file_path); time.sleep(1); size2 = os.path.getsize(zip_file_path)
                    if size1 == size2:
                        print(f"[BV] ZIP détecté : {os.path.basename(zip_file_path)}", file=self.stdout_redirect)
                        break
                time.sleep(1)

            if not zip_file_path:
                print("[BV] Aucun fichier ZIP trouvé.", file=self.stdout_redirect)
            else:
                if os.path.exists(target_path):
                    print(f"[BV] Remplacement du dossier '{target_folder_name}'", file=self.stdout_redirect)
                    shutil.rmtree(target_path, ignore_errors=True)
                os.makedirs(target_path, exist_ok=True)
                with zipfile.ZipFile(zip_file_path, 'r') as zf:
                    zf.extractall(path=target_path)
                os.remove(zip_file_path)
                print(f"[BV] Décompression terminée dans '{target_folder_name}'", file=self.stdout_redirect)
                self._set_status(f"Bassin versant : {target_path}")
        except Exception as e:
            print(f"[BV] Erreur décompression : {e}", file=self.stdout_redirect)
        finally:
            self.after(0, lambda: self.bassin_btn.config(state="normal"))

    def _detect_commune(self, lat: float, lon: float) -> str:
        try:
            url = (f"https://nominatim.openstreetmap.org/reverse?format=json"
                   f"&lat={lat}&lon={lon}&zoom=10&addressdetails=1")
            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=10) as resp:
                data = json.load(resp)
            addr = data.get("address", {})
            for key in ("city", "town", "village", "municipality"):
                if key in addr:
                    return addr[key]
        except Exception as e:
            print(f"[IGN] Détection commune échouée : {e}", file=self.stdout_redirect)
        return "Inconnue"

    def _start_thread(self):
        if not self.coord_var.get().strip():
            messagebox.showerror("Erreur", "Renseigner les coordonnées en DMS."); return
        self.run_btn.config(state="disabled")
        t = threading.Thread(target=self._run_process)
        t.daemon = True
        t.start()

    def _run_process(self):
        try:
            # Sauvegarde préférences
            self.prefs.update({
                "RLT_COORD": self.coord_var.get().strip(),
                "RLT_WAIT": float(self.wait_var.get()),
                "RLT_OUT": self.out_dir_var.get().strip(),
                "RLT_HEADLESS": bool(self.headless_var.get()),
            }); save_prefs(self.prefs)

            print(f"[IGN] Parsing coordonnées…", file=self.stdout_redirect)
            # La chaîne doit contenir LAT puis LON en DMS, séparés par espace/virgule/tab
            parts = re.split(r"\s{2,}|,|\t", self.coord_var.get().strip())
            if len(parts) < 2:
                # fallback : tente un split sur premier espace
                parts = re.split(r"\s+", self.coord_var.get().strip(), maxsplit=1)
            if len(parts) < 2:
                raise ValueError("Coordonnées DMS attendues au format « LAT  LON »")

            lat_dd = dms_to_dd(parts[0])
            lon_dd = dms_to_dd(parts[1])
            commune = self._detect_commune(lat_dd, lon_dd)
            self.after(0, lambda: self.commune_var.set(commune))
            print(f"[IGN] Commune détectée : {commune}", file=self.stdout_redirect)
            wait_s = float(self.wait_var.get())
            out_dir = self.out_dir_var.get().strip() or OUTPUT_DIR_RLT
            os.makedirs(out_dir, exist_ok=True)
            comment_txt = COMMENT_TEMPLATE.format(commune=commune)

            drv_opts = webdriver.ChromeOptions()
            drv_opts.add_argument("--log-level=3")
            drv_opts.add_experimental_option('excludeSwitches', ['enable-logging'])
            drv_opts.add_argument("--disable-extensions")
            if self.headless_var.get():
                drv_opts.add_argument("--headless=new")

            print(f"[IGN] Lancement Chrome…", file=self.stdout_redirect)
            driver = webdriver.Chrome(options=drv_opts)
            try:
                driver.maximize_window()
            except Exception:
                pass

            images = []
            viewport = (By.CSS_SELECTOR, "div.ol-viewport")

            for title, layer_val in LAYERS:
                url = URL.format(lon=f"{lon_dd:.6f}", lat=f"{lat_dd:.6f}", layer=layer_val)
                print(f"[IGN] {title} → {url}", file=self.stdout_redirect)
                driver.get(url)
                WebDriverWait(driver, 20).until(EC.visibility_of_element_located(viewport))
                time.sleep(wait_s)
                tgt = driver.find_element(*viewport)
                img_path = os.path.join(out_dir, f"{title}.png")
                if tgt.screenshot(img_path):
                    img = Image.open(img_path)
                    w, h = img.size
                    left, right = int(w * 0.05), int(w * 0.95)
                    img.crop((left, 0, right, h)).save(img_path)
                    images.append((title, img_path))
                    print(f"[IGN] Capture OK : {img_path}", file=self.stdout_redirect)
                else:
                    print(f"[IGN] Capture échouée : {title}", file=self.stdout_redirect)

            driver.quit()

            if not images:
                print("[IGN] Aucune image → pas de doc.", file=self.stdout_redirect)
                messagebox.showwarning("IGN", "Aucune image capturée.")
                return

            print("[IGN] Génération du Word…", file=self.stdout_redirect)
            doc = Document()
            style_normal = doc.styles['Normal']
            style_normal.font.name = 'Calibri'
            style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            style_tbl = doc.styles['Table Grid']
            style_tbl.font.name = 'Calibri'
            style_tbl._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

            sec = doc.sections[0]
            sec.orientation = WD_ORIENT.LANDSCAPE
            sec.page_width, sec.page_height = sec.page_height, sec.page_width
            for m in (sec.left_margin, sec.right_margin, sec.top_margin, sec.bottom_margin):
                m = Cm(1.5)

            cap_par = doc.add_paragraph()
            cap_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_hyperlink(cap_par, "https://remonterletemps.ign.fr/",
                          f"Comparaison temporelle — {commune} (source : IGN – RemonterLeTemps)")

            table = doc.add_table(rows=2, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            table.autofit = False

            for idx, (title, path) in enumerate(images):
                r, c = divmod(idx, 2)
                cell = table.cell(r, c)
                p_t = cell.paragraphs[0]
                run_t = p_t.add_run(title); run_t.bold = True
                p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.add_paragraph()
                p_img = cell.add_paragraph()
                if os.path.exists(path):
                    p_img.add_run().add_picture(path, width=IMG_WIDTH)
                else:
                    p_img.add_run(f"[image manquante : {title}]")
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()
            p_comm = doc.add_paragraph(comment_txt)
            p_comm.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            doc_path = os.path.join(out_dir, WORD_FILENAME)
            doc.save(doc_path)
            print(f"[IGN] Document généré : {doc_path}", file=self.stdout_redirect)
            self._set_status(f"Terminé — {doc_path}")
        except Exception as e:
            print(f"[IGN] Erreur : {e}", file=self.stdout_redirect)
            messagebox.showerror("IGN", str(e))
        finally:
            self.after(0, lambda: self.run_btn.config(state="normal"))

    def _set_status(self, txt: str):
        self.after(0, lambda: self.status_label.config(text=txt))

