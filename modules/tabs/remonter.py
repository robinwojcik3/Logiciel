import os
import re
import time
import threading
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter import font as tkfont

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from selenium.webdriver import ActionChains

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from PIL import Image

from ..utils import TextRedirector, ToolTip, save_prefs, log_with_time

LAYERS = [
    ("Aujourd’hui",   "10"),
    ("2000-2005",     "18"),
    ("1965-1980",     "20"),
    ("1950-1965",     "19"),
]
URL = ("https://remonterletemps.ign.fr/comparer/?lon={lon}&lat={lat}"
       "&z=17&layer1={layer}&layer2=19&mode=dub1")
WAIT_TILES_DEFAULT = 1.5
IMG_WIDTH = Cm(12.5 * 0.8)
WORD_FILENAME = "Comparaison_temporelle_Paysage.docx"
OUTPUT_DIR_RLT = os.path.join("C:\\Users\\utilisateur\\Mon Drive\\1 - Bota & Travail\\+++++++++  BOTA  +++++++++\\---------------------- 3) BDD\\PYTHON\\2) Contexte éco\\OUTPUT", "Remonter le temps")
COMMENT_TEMPLATE = (
    "Rédige un commentaire synthétique de l'évolution de l'occupation du sol observée "
    "sur les images aériennes de la zone d'étude, aux différentes dates indiquées "
    "(1950–1965, 1965–1980, 2000–2005, aujourd’hui). Concentre-toi sur les grandes "
    "dynamiques d'aménagement (urbanisation, artificialisation, évolution des milieux "
    "ouverts ou boisés), en identifiant les principales transformations visibles. "
    "Fais ta réponse en un seul court paragraphe. Intègre les éléments de contexte "
    "historique et territorial propres à la commune de {commune} pour interpréter ces évolutions."
)


def dms_to_dd(text: str) -> float:
    pat = r"(\d{1,3})[°d]\s*(\d{1,2})['m]\s*([\d\.]+)[\"s]?\s*([NSEW])"
    alt = r"(\d{1,3})\s+(\d{1,2})\s+([\d\.]+)\s*([NSEW])"
    m = re.search(pat, text, re.I) or re.search(alt, text, re.I)
    if not m:
        raise ValueError(f"Format DMS invalide : {text}")
    deg, mn, sc, hemi = m.groups()
    dd = float(deg) + float(mn)/60 + float(sc)/3600
    return -dd if hemi.upper() in ("S", "W") else dd


def add_hyperlink(paragraph, url: str, text: str, italic: bool = True):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    from docx.oxml import OxmlElement, ns
    fld_simple = OxmlElement('w:hyperlink')
    fld_simple.set(ns.qn('r:id'), r_id)

    run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    if italic:
        i = OxmlElement('w:i')
        r_pr.append(i)
    u = OxmlElement('w:u')
    u.set(ns.qn('w:val'), 'single')
    r_pr.append(u)
    run.append(r_pr)
    run_text = OxmlElement('w:t')
    run_text.text = text
    run.append(run_text)
    fld_simple.append(run)
    paragraph._p.append(fld_simple)


class RemonterLeTempsTab(ttk.Frame):
    def __init__(self, parent, style_helper, prefs: dict):
        super().__init__(parent, padding=12)
        self.style_helper = style_helper
        self.prefs = prefs

        self.font_title = tkfont.Font(family="Segoe UI", size=15, weight="bold")
        self.font_sub = tkfont.Font(family="Segoe UI", size=10)
        self.font_mono = tkfont.Font(family="Consolas", size=9)

        self.coord_var = tk.StringVar(value=self.prefs.get("RLT_COORD", ""))
        self.commune_var = tk.StringVar()
        self.wait_var = tk.DoubleVar(value=float(self.prefs.get("RLT_WAIT", WAIT_TILES_DEFAULT)))
        self.out_dir_var = tk.StringVar(value=self.prefs.get("RLT_OUT", OUTPUT_DIR_RLT))
        self.headless_var = tk.BooleanVar(value=bool(self.prefs.get("RLT_HEADLESS", False)))

        self._build_ui()

    def _build_ui(self):
        header = ttk.Frame(self, style="Header.TFrame", padding=(14, 12))
        header.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(header, text="Remonter le temps & Bassin versant", style="Card.TLabel", font=self.font_title)\
            .grid(row=0, column=0, sticky="w")
        ttk.Label(header, text="Capture des vues IGN et génération d'un Word.", style="Subtle.TLabel", font=self.font_sub)\
            .grid(row=1, column=0, sticky="w", pady=(4,0))
        header.columnconfigure(0, weight=1)

        card = ttk.Frame(self, style="Card.TFrame", padding=12)
        card.pack(fill=tk.X)
        ttk.Label(card, text="Coordonnées (DMS)", style="Card.TLabel").grid(row=0, column=0, sticky="w")
        ttk.Entry(card, textvariable=self.coord_var).grid(row=0, column=1, sticky="ew")
        card.columnconfigure(1, weight=1)

        ttk.Label(card, text="Pause chargement (s)", style="Card.TLabel").grid(row=1, column=0, sticky="w", pady=(8,0))
        ttk.Entry(card, textvariable=self.wait_var).grid(row=1, column=1, sticky="w")

        ttk.Label(card, text="Dossier de sortie", style="Card.TLabel").grid(row=2, column=0, sticky="w", pady=(8,0))
        row = ttk.Frame(card, style="Card.TFrame")
        row.grid(row=2, column=1, sticky="ew")
        row.columnconfigure(0, weight=1)
        ttk.Entry(row, textvariable=self.out_dir_var).grid(row=0, column=0, sticky="ew")
        ttk.Button(row, text="Parcourir…", command=self._pick_out_dir).grid(row=0, column=1, padx=(6,0))

        ttk.Checkbutton(card, text="Mode headless", variable=self.headless_var, style="Card.TCheckbutton")\
            .grid(row=3, column=0, columnspan=2, sticky="w", pady=(8,0))

        act = ttk.Frame(self, style="Card.TFrame", padding=12)
        act.pack(fill=tk.X, pady=(10,0))
        self.run_btn = ttk.Button(act, text="▶ Lancer", style="Accent.TButton", command=self._start_thread)
        self.run_btn.grid(row=0, column=0, sticky="w")
        self.status_label = ttk.Label(act, text="", style="Status.TLabel")
        self.status_label.grid(row=0, column=1, sticky="e")
        act.columnconfigure(1, weight=1)

        bottom = ttk.Frame(self, style="Card.TFrame", padding=12)
        bottom.pack(fill=tk.BOTH, expand=True, pady=(10,0))
        self.log_text = tk.Text(bottom, height=12, wrap=tk.WORD, state='disabled',
                                bg=self.style_helper.style.lookup("Card.TFrame", "background"),
                                fg=self.style_helper.style.lookup("TLabel", "foreground"))
        self.log_text.configure(font=self.font_mono, relief="flat")
        log_scroll = ttk.Scrollbar(bottom, orient="vertical", command=self.log_text.yview)
        self.log_text['yscrollcommand'] = log_scroll.set
        log_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.stdout_redirect = TextRedirector(self.log_text)

    def _pick_out_dir(self):
        base = self.out_dir_var.get() or os.path.expanduser("~")
        d = filedialog.askdirectory(title="Choisir le dossier de sortie",
                                    initialdir=base if os.path.isdir(base) else os.path.expanduser("~"))
        if d:
            self.out_dir_var.set(d)

    def _start_thread(self):
        self.run_btn.config(state="disabled")
        t = threading.Thread(target=self._run_process)
        t.daemon = True
        t.start()

    def _detect_commune(self, lat_dd: float, lon_dd: float) -> str:
        url = f"https://nominatim.openstreetmap.org/reverse?format=json&lat={lat_dd}&lon={lon_dd}&zoom=10&addressdetails=1"
        try:
            import urllib.request, json
            with urllib.request.urlopen(url) as resp:
                data = json.load(resp)
                return data['address'].get('town') or data['address'].get('city') or ''
        except Exception:
            return ''

    def _run_process(self):
        try:
            if len(self.coord_var.get().split()) < 2:
                parts = re.split(r"\s+", self.coord_var.get().strip())
            else:
                parts = self.coord_var.get().split()
            if len(parts) < 2:
                raise ValueError("Coordonnées DMS attendues au format « LAT  LON »")

            lat_dd = dms_to_dd(parts[0])
            lon_dd = dms_to_dd(parts[1])
            commune = self._detect_commune(lat_dd, lon_dd)
            self.after(0, lambda: self.commune_var.set(commune))
            print(f"[IGN] Commune détectée : {commune}", file=self.stdout_redirect)
            wait_s = float(self.wait_var.get())
            out_dir = self.out_dir_var.get().strip() or OUTPUT_DIR_RLT
            os.makedirs(out_dir, exist_ok=True)
            comment_txt = COMMENT_TEMPLATE.format(commune=commune)

            drv_opts = webdriver.ChromeOptions()
            drv_opts.add_argument("--log-level=3")
            drv_opts.add_experimental_option('excludeSwitches', ['enable-logging'])
            drv_opts.add_argument("--disable-extensions")
            if self.headless_var.get():
                drv_opts.add_argument("--headless=new")

            print(f"[IGN] Lancement Chrome…", file=self.stdout_redirect)
            driver = webdriver.Chrome(options=drv_opts)
            try:
                driver.maximize_window()
            except Exception:
                pass

            images = []
            viewport = (By.CSS_SELECTOR, "div.ol-viewport")

            for title, layer_val in LAYERS:
                url = URL.format(lon=f"{lon_dd:.6f}", lat=f"{lat_dd:.6f}", layer=layer_val)
                print(f"[IGN] {title} → {url}", file=self.stdout_redirect)
                driver.get(url)
                WebDriverWait(driver, 20).until(EC.visibility_of_element_located(viewport))
                time.sleep(wait_s)
                tgt = driver.find_element(*viewport)
                img_path = os.path.join(out_dir, f"{title}.png")
                if tgt.screenshot(img_path):
                    img = Image.open(img_path)
                    w, h = img.size
                    left, right = int(w * 0.05), int(w * 0.95)
                    img.crop((left, 0, right, h)).save(img_path)
                    images.append((title, img_path))
                    print(f"[IGN] Capture OK : {img_path}", file=self.stdout_redirect)
                else:
                    print(f"[IGN] Capture échouée : {title}", file=self.stdout_redirect)

            driver.quit()

            if not images:
                print("[IGN] Aucune image → pas de doc.", file=self.stdout_redirect)
                messagebox.showwarning("IGN", "Aucune image capturée.")
                return

            print("[IGN] Génération du Word…", file=self.stdout_redirect)
            doc = Document()
            style_normal = doc.styles['Normal']
            style_normal.font.name = 'Calibri'
            style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            style_tbl = doc.styles['Table Grid']
            style_tbl.font.name = 'Calibri'
            style_tbl._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

            sec = doc.sections[0]
            sec.orientation = WD_ORIENT.LANDSCAPE
            sec.page_width, sec.page_height = sec.page_height, sec.page_width
            for m in (sec.left_margin, sec.right_margin, sec.top_margin, sec.bottom_margin):
                m = Cm(1.5)

            cap_par = doc.add_paragraph()
            cap_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_hyperlink(cap_par, "https://remonterletemps.ign.fr/",
                          f"Comparaison temporelle — {commune} (source : IGN – RemonterLeTemps)")

            table = doc.add_table(rows=2, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            table.autofit = False

            for idx, (title, path) in enumerate(images):
                r, c = divmod(idx, 2)
                cell = table.cell(r, c)
                p_t = cell.paragraphs[0]
                run_t = p_t.add_run(title)
                run_t.bold = True
                p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.add_paragraph()
                p_img = cell.add_paragraph()
                if os.path.exists(path):
                    p_img.add_run().add_picture(path, width=IMG_WIDTH)
                else:
                    p_img.add_run(f"[image manquante : {title}]")
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()
            p_comm = doc.add_paragraph(comment_txt)
            p_comm.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            doc_path = os.path.join(out_dir, WORD_FILENAME)
            doc.save(doc_path)
            print(f"[IGN] Document généré : {doc_path}", file=self.stdout_redirect)
            self._set_status(f"Terminé — {doc_path}")
        except Exception as e:
            print(f"[IGN] Erreur : {e}", file=self.stdout_redirect)
            messagebox.showerror("IGN", str(e))
        finally:
            self.after(0, lambda: self.run_btn.config(state="normal"))

    def _set_status(self, txt: str):
        self.after(0, lambda: self.status_label.config(text=txt))
