"""Generate Word report for ecological context by inserting tables and maps.

This module combines the results of the "ID contexte éco" analysis and the
exported maps into a single Word document based on a template. It expects an
Excel file produced by :mod:`id_contexte_eco` and a folder containing PNG maps
exported from QGIS.

Example
-------
>>> from modules.report_generator import generate_contexte_eco_report
>>> generate_contexte_eco_report(
...     excel_path="output/contexte_eco.xlsx",
...     maps_dir="Cartes contexte éco export",
...     template_path="Template word  Contexte éco/1 Template Contexte éco.docx",
...     output_path="output/rapport_contexte_eco.docx",
... )
"""

from __future__ import annotations

import os
import shutil
from typing import Dict, Mapping, Optional

import pandas as pd
from docx import Document
from docx.shared import Cm

# Configuration par défaut pour les différents zonages.
# Chaque entrée indique le nom de l'onglet Excel et le fichier image correspondant.
DEFAULT_ZONAGES: Mapping[str, Dict[str, str]] = {
    "NATURA2000": {
        "sheet": "Natura 2000",
        "image": "Contexte éco - N2000__AE.png",
    },
    "ENS": {
        "sheet": "ENS",
        "image": "Contexte éco - ENS__AE.png",
    },
    "APPB": {
        "sheet": "APPB",
        "image": "Contexte éco - APPB__AE.png",
    },
}


def _replace_with_table(paragraph, df: pd.DataFrame) -> None:
    """Insère un tableau à la place du paragraphe fourni."""
    tbl = paragraph._parent.add_table(rows=1, cols=len(df.columns))
    hdr_cells = tbl.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr_cells[j].text = str(col)
    for _, row in df.iterrows():
        row_cells = tbl.add_row().cells
        for j, col in enumerate(df.columns):
            row_cells[j].text = str(row[col])
    paragraph._p.addnext(tbl._element)
    paragraph._element.getparent().remove(paragraph._element)


def _replace_with_image(paragraph, image_path: str, width_cm: float = 15) -> None:
    """Insère une image à la place du paragraphe fourni."""
    paragraph.text = ""
    run = paragraph.add_run()
    run.add_picture(image_path, width=Cm(width_cm))


def generate_contexte_eco_report(
    excel_path: str,
    maps_dir: str,
    template_path: str,
    output_path: str,
    zonages: Optional[Mapping[str, Dict[str, str]]] = None,
) -> str:
    """Crée un rapport Word à partir d'un modèle.

    Parameters
    ----------
    excel_path: str
        Chemin vers le fichier Excel généré par l'analyse « ID contexte éco ».
    maps_dir: str
        Dossier contenant les cartes PNG exportées depuis QGIS.
    template_path: str
        Chemin du fichier Word modèle à copier.
    output_path: str
        Chemin du document Word généré.
    zonages: Mapping, optional
        Dictionnaire configurant les onglets Excel et les fichiers images pour
        chaque type de zonage. Par défaut : ``DEFAULT_ZONAGES``.

    Returns
    -------
    str
        Chemin du rapport Word généré.
    """
    zonages = zonages or DEFAULT_ZONAGES

    # Toujours travailler sur une copie du modèle.
    shutil.copy(template_path, output_path)
    doc = Document(output_path)

    for key, cfg in zonages.items():
        sheet_name = cfg.get("sheet")
        img_name = cfg.get("image")

        table_marker = f"TABLEAU {key}"
        image_marker = f"CARTE {key}"

        if sheet_name:
            df = pd.read_excel(excel_path, sheet_name=sheet_name)
            for p in doc.paragraphs:
                if table_marker in p.text:
                    _replace_with_table(p, df)
                    break

        if img_name:
            img_path = os.path.join(maps_dir, img_name)
            if os.path.exists(img_path):
                for p in doc.paragraphs:
                    if image_marker in p.text:
                        _replace_with_image(p, img_path)
                        break

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Génère un rapport Contexte éco")
    parser.add_argument("excel")
    parser.add_argument("maps")
    parser.add_argument("template")
    parser.add_argument("output")
    args = parser.parse_args()

    generate_contexte_eco_report(args.excel, args.maps, args.template, args.output)
