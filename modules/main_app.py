#!/usr/bin/env "C:/Program Files/QGIS 3.40.3/apps/Python312/python.exe"

# -*- coding: utf-8 -*-

r"""

Application à onglets (ttk.Notebook) :



Onglet « Contexte éco » :

    - Export des mises en page QGIS en PNG.

    - Identification des zonages (ID Contexte éco) avec tampon configurable.

    - Sélection commune des shapefiles ZE/AE et console partagée.

    - Boutons « Remonter le temps », « Ouvrir Google Maps » et « Bassin versant »

      utilisant le centroïde de la zone d'étude.



Onglet « Identification Pl@ntNet » :

    - Reconnaissance de plantes via l'API Pl@ntNet.



Pré-requis Python : qgis (environnement QGIS), selenium, pillow, python-docx,

openpyxl (non utilisé ici), chromedriver dans PATH.

"""



import os

import sys

import re

import json

import time

import shutil

import tempfile

import datetime

import threading

import urllib.request

import webbrowser

import tkinter as tk

from tkinter import ttk, filedialog, messagebox

from tkinter import font as tkfont

from typing import List, Optional, Tuple

from concurrent.futures import ProcessPoolExecutor, ThreadPoolExecutor, as_completed

import requests

from io import BytesIO

import pillow_heif

import zipfile

import traceback

import subprocess

import geopandas as gpd

from bs4 import BeautifulSoup



# ==== Imports supplémentaires pour l'onglet Contexte éco ====

# Note: geopandas n'est pas utilisé directement dans ce module.

# Les traitements géospatiaux sont effectués dans des modules dédiés

# (ex: id_contexte_eco) afin d'éviter de charger des dépendances lourdes

# au lancement de l'UI principale.



# Assurer que le dossier racine du projet est dans sys.path quand ce fichier est exécuté directement
try:
    _THIS_DIR = os.path.dirname(__file__)
    _PROJECT_ROOT = os.path.abspath(os.path.join(_THIS_DIR, '..'))  # .. = dossier Logiciel/
    if _PROJECT_ROOT not in sys.path:
        sys.path.insert(0, _PROJECT_ROOT)
except Exception:
    pass


# Import du scraper Wikipédia

try:
    # When run as package (via Start.py)
    from .wikipedia_scraper import DEP, get_wikipedia_extracts
except Exception:
    # Fallback when running this file directly
    from modules.wikipedia_scraper import DEP, get_wikipedia_extracts
  
  
  # Import du worker QGIS externalisé
 
try:
    from .export_worker import worker_run
except Exception:
    from modules.export_worker import worker_run
  
  
  
  
  

# ==== Imports spécifiques onglet 2 (gardés en tête de fichier comme le script source) ====

from selenium import webdriver

from selenium.webdriver.chrome.service import Service

from selenium.webdriver.common.by import By

from selenium.webdriver.support.ui import WebDriverWait

from selenium.webdriver.support import expected_conditions as EC

from selenium.webdriver.common.keys import Keys

from selenium.webdriver import ActionChains



from docx import Document

from docx.shared import Cm

from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.enum.section import WD_ORIENT

from docx.enum.table import WD_TABLE_ALIGNMENT

from docx.oxml.ns import qn
from docx.oxml import OxmlElement



from PIL import Image

from bs4 import BeautifulSoup



# Enregistrer le décodeur HEIF

pillow_heif.register_heif_opener()



# =========================

# Paramètres globaux

# =========================

# Contexte éco — Export Cartes

DPI_DEFAULT        = 300

N_WORKERS_DEFAULT  = max(1, min((os.cpu_count() or 2) - 1, 6))

MARGIN_FAC_DEFAULT = 1.15

OVERWRITE_DEFAULT  = False



LAYER_AE_NAME = "Aire d'étude élargie"

LAYER_ZE_NAME = "Zone d'étude"



BASE_SHARE = r"\\192.168.1.240\commun\PARTAGE"

SUBPATH    = r"Espace_RWO\CARTO ROBIN"



OUT_IMG    = r"C:\Users\utilisateur\Mon Drive\1 - Bota & Travail\+++++++++  BOTA  +++++++++\---------------------- 3) BDD\PYTHON\2) Contexte éco\OUTPUT"



# Dossier par défaut pour la sélection des shapefiles (onglet 1)

DEFAULT_SHAPE_DIR = r"C:\Users\utilisateur\Mon Drive\1 - Bota & Travail\+++++++++  BOTA  +++++++++\---------------------- 2) CARTO terrain"



# QGIS

QGIS_ROOT = r"C:\Program Files\QGIS 3.40.3"

QGIS_APP  = os.path.join(QGIS_ROOT, "apps", "qgis")

PY_VER    = "Python312"



# Préférences

PREFS_PATH = os.path.join(os.path.expanduser("~"), "ExportCartesContexteEco.config.json")



# Constantes « Remonter le temps » et « Bassin versant » (issues du script source)

LAYERS = [

    ("Aujourd’hui",   "10"),

    ("2000-2005",     "18"),

    ("1965-1980",     "20"),

    ("1950-1965",     "19"),

]

URL = ("https://remonterletemps.ign.fr/comparer/?lon={lon}&lat={lat}"

       "&z=17&layer1={layer}&layer2=19&mode=dub1")

WAIT_TILES_DEFAULT = 1.5

IMG_WIDTH = Cm(12.5 * 0.8)

WORD_FILENAME = "Comparaison_temporelle_Paysage.docx"

OUTPUT_DIR_RLT = os.path.join(OUT_IMG, "Remonter le temps")

COMMENT_TEMPLATE = (

    "Rédige un commentaire synthétique de l'évolution de l'occupation du sol observée "

    "sur les images aériennes de la zone d'étude, aux différentes dates indiquées "

    "(1950–1965, 1965–1980, 2000–2005, aujourd’hui). Concentre-toi sur les grandes "

    "dynamiques d'aménagement (urbanisation, artificialisation, évolution des milieux "

    "ouverts ou boisés), en identifiant les principales transformations visibles. "

    "Fais ta réponse en un seul court paragraphe. Intègre les éléments de contexte "

    "historique et territorial propres à la commune de {commune} pour interpréter ces évolutions."

)



# Onglet 3 — Identification Pl@ntNet

API_KEY = "2b10vfT6MvFC2lcAzqG1ZMKO"  # Votre clé API Pl@ntNet

PROJECT = "all"

API_URL = f"https://my-api.plantnet.org/v2/identify/{PROJECT}?api-key={API_KEY}"



# =========================

# Utils communs

# =========================

def log_with_time(msg: str) -> None:

    print(f"[{datetime.datetime.now().strftime('%H:%M:%S')}] {msg}")



def normalize_name(s: str) -> str:

    s2 = s.replace("\u00A0", " ").replace("\u202F", " ")

    while "  " in s2:

        s2 = s2.replace("  ", " ")

    return s2.strip().lower()



def to_long_unc(path: str) -> str:

    if path.startswith("\\\\?\\"): return path

    if path.startswith("\\\\"):   return "\\\\?\\UNC" + path[1:]

    return "\\\\?\\" + path


def from_long_unc(path: str) -> str:

    r"""Convertit un chemin Windows étendu (\\?\ ou \\?\UNC) en chemin standard.

    - \\?\UNC\server\share\... -> \\server\share\...
    - \\?\C:\path -> C:\path
    """

    p = path or ""

    if p.startswith("\\\\?\\UNC"):
        return "\\\\" + p[8:]

    if p.startswith("\\\\?\\"):
        return p[4:]

    return p



def chunk_even(lst: List[str], k: int) -> List[List[str]]:

    if not lst: return []

    k = max(1, min(k, len(lst)))

    base = len(lst) // k

    extra = len(lst) % k

    out, start = [], 0

    for i in range(k):

        size = base + (1 if i < extra else 0)

        out.append(lst[start:start+size])

        start += size

    return out



def qgis_multiprocessing_ok() -> bool:

    """Quick check: can QGIS's Python import _multiprocessing? Avoids noisy spawn failures."""

    try:

        qgis_py = os.path.join(QGIS_ROOT, "apps", PY_VER, "python.exe")

        if not os.path.isfile(qgis_py):

            return False

        env = os.environ.copy()

        # Prefer clean env and explicit paths

        qgis_py_root = os.path.join(QGIS_ROOT, "apps", PY_VER)

        qgis_lib = os.path.join(qgis_py_root, "Lib")

        qgis_dlls = os.path.join(qgis_py_root, "DLLs")

        qgis_site = os.path.join(qgis_lib, "site-packages")

        qgis_app_py = os.path.join(QGIS_APP, "python")

        env.pop("PYTHONHOME", None); env.pop("PYTHONPATH", None); env["PYTHONNOUSERSITE"] = "1"

        env["PYTHONHOME"] = qgis_py_root

        env["PYTHONPATH"] = os.pathsep.join([qgis_py_root, qgis_lib, qgis_dlls, qgis_site, qgis_app_py])

        code = "import multiprocessing, multiprocessing.connection, _multiprocessing; print('OK')"

        r = subprocess.run([qgis_py, "-c", code], stdout=subprocess.PIPE, stderr=subprocess.STDOUT,

                           text=True, timeout=6, env=env, cwd=os.path.dirname(__file__))

        return r.returncode == 0 and (r.stdout or "").strip().endswith("OK")

    except Exception:

        return False



def run_worker_subprocess(projects: List[str], cfg: dict) -> tuple[int, int]:

    """Exécute un lot via QGIS Python dans un sous-processus.



    Cette approche évite totalement multiprocessing dans l'interpréteur QGIS,

    supprimant l'erreur `_multiprocessing`.

    """

    import json as _json

    tmp = None

    try:

        qgis_py = os.path.join(QGIS_ROOT, "apps", PY_VER, "python.exe")

        if not os.path.isfile(qgis_py):

            raise RuntimeError(f"Python QGIS introuvable: {qgis_py}")



        repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir))

        qgis_py_root = os.path.join(QGIS_ROOT, "apps", PY_VER)

        qgis_lib  = os.path.join(qgis_py_root, "Lib")

        qgis_dlls = os.path.join(qgis_py_root, "DLLs")

        qgis_site = os.path.join(qgis_lib, "site-packages")

        qgis_app_py = os.path.join(QGIS_APP, "python")



        data = {"projects": list(projects or []), "cfg": dict(cfg or {})}

        fd, tmp = tempfile.mkstemp(prefix="qgis_worker_", suffix=".json")

        os.close(fd)

        with open(tmp, "w", encoding="utf-8") as f:

            _json.dump(data, f)



        env = os.environ.copy()

        for k in ("PYTHONPATH", "PYTHONHOME", "PYTHONSTARTUP"):

            env.pop(k, None)

        env["PYTHONNOUSERSITE"] = "1"

        env["PYTHONHOME"] = qgis_py_root

        env["PYTHONPATH"] = os.pathsep.join([repo_root, qgis_py_root, qgis_lib, qgis_dlls, qgis_site, qgis_app_py])



        cmd = [qgis_py, "-m", "modules.export_worker_cli", tmp]

        r = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, env=env)

        out = (r.stdout or "").strip()

        if r.returncode == 0 and "," in out:

            try:

                last = out.splitlines()[-1]

                s_ok, s_ko = last.split(",", 1)

                return int(s_ok), int(s_ko)

            except Exception:

                pass

        log_with_time(f"Worker subproc KO (code={r.returncode}): {out[:400]}")

        return 0, len(projects or [])

    finally:

        if tmp and os.path.isfile(tmp):

            try:

                os.remove(tmp)

            except Exception:

                pass



def load_prefs() -> dict:

    if os.path.isfile(PREFS_PATH):

        try:

            with open(PREFS_PATH, "r", encoding="utf-8") as f:

                return json.load(f)

        except Exception:

            return {}

    return {}



def save_prefs(d: dict) -> None:

    try:

        with open(PREFS_PATH, "w", encoding="utf-8") as f:

            json.dump(d, f, ensure_ascii=False, indent=2)

    except Exception:

        pass



class TextRedirector:

    def __init__(self, widget): self.widget = widget

    def write(self, s):

        self.widget.config(state='normal')

        self.widget.insert(tk.END, s)

        self.widget.see(tk.END)

        self.widget.config(state='disabled')

        self.widget.update_idletasks()

    def flush(self): pass



class ToolTip:

    def __init__(self, widget, text: str, delay: int = 600):

        self.widget = widget; self.text = text; self.delay = delay

        self.tipwindow = None; self.id = None

        widget.bind("<Enter>", self._schedule); widget.bind("<Leave>", self._hide)

    def _schedule(self, _=None):

        self._cancel(); self.id = self.widget.after(self.delay, self._show)

    def _cancel(self):

        if self.id: self.widget.after_cancel(self.id); self.id = None

    def _show(self):

        if self.tipwindow: return

        x = self.widget.winfo_rootx() + 15; y = self.widget.winfo_rooty() + 25

        self.tipwindow = tw = tk.Toplevel(self.widget); tw.wm_overrideredirect(True); tw.wm_geometry(f"+{x}+{y}")

        ttk.Label(tw, text=self.text, style="Tooltip.TLabel", padding=(8, 4)).pack()

    def _hide(self, _=None):

        self._cancel()

        if self.tipwindow: self.tipwindow.destroy(); self.tipwindow = None



# =========================

# Fonctions Pl@ntNet

# =========================

def resize_image(image_path, max_size=(800, 800), quality=70):

    """

    Redimensionne et compresse une image.



    :param image_path: Chemin de l'image à traiter.

    :param max_size: Tuple indiquant la taille maximale (largeur, hauteur).

    :param quality: Qualité de compression (1-100).

    :return: BytesIO de l'image traitée ou None en cas d'erreur.

    """

    try:

        with Image.open(image_path) as img:

            img.thumbnail(max_size)

            buffer = BytesIO()

            img.save(buffer, format='JPEG', quality=quality)

            buffer.seek(0)

            return buffer

    except Exception as e:

        print(f"Erreur lors du redimensionnement de l'image : {e}")

        return None



def identify_plant(image_path, organ):

    """

    Envoie une image à l'API Pl@ntNet pour identification.



    :param image_path: Chemin de l'image à envoyer.

    :param organ: Type d'organe de la plante (par exemple, 'flower').

    :return: Nom scientifique de la plante identifiée ou None.

    """

    print(f"Envoi de l'image à l'API : {image_path}")

    try:

        resized_image = resize_image(image_path)

        if not resized_image:

            print(f"Échec du redimensionnement de l'image : {image_path}")

            return None



        files = {

            'images': (os.path.basename(image_path), resized_image, 'image/jpeg')

        }

        data = {

            'organs': organ

        }



        response = requests.post(API_URL, files=files, data=data)



        print(f"Réponse de l'API : {response.status_code}")

        if response.status_code == 200:

            json_result = response.json()

            try:

                species = json_result['results'][0]['species']['scientificNameWithoutAuthor']

                print(f"Plante identifiée : {species}")

                return species

            except (KeyError, IndexError):

                print(f"Aucun résultat trouvé pour l'image : {image_path}")

                return None

        else:

            print(f"Erreur API : {response.status_code} - {response.text}")

            return None

    except Exception as e:

        print(f"Exception lors de l'identification de la plante : {e}")

        return None



def copy_and_rename_file(file_path, dest_folder, new_name, count):

    """

    Copie et renomme un fichier dans le dossier de destination.



    :param file_path: Chemin du fichier original.

    :param dest_folder: Dossier de destination.

    :param new_name: Nom scientifique de la plante.

    :param count: Compteur pour différencier les fichiers portant le même nom.

    """

    ext = os.path.splitext(file_path)[1]

    if count == 1:

        new_file_name = f"{new_name} @plantnet{ext}"

    else:

        new_file_name = f"{new_name} @plantnet({count}){ext}"

    new_path = os.path.join(dest_folder, new_file_name)
    
    try:
        # Ensure destination directory exists
        os.makedirs(dest_folder, exist_ok=True)

        shutil.copy(file_path, new_path)

        print(f"Fichier copié et renommé : {file_path} -> {new_path}")

    except Exception as e:

        print(f"Erreur lors de la copie du fichier : {e}")





def discover_projects() -> List[str]:

    partage = BASE_SHARE

    base_dir: Optional[str] = None

    try:

        if os.path.isdir(partage):

            for e in os.listdir(partage):

                if normalize_name(e) == normalize_name("Espace_RWO"):

                    base_espace = os.path.join(partage, e)

                    for s in os.listdir(base_espace):

                        if normalize_name(s) == normalize_name("CARTO ROBIN"):

                            base_dir = os.path.join(base_espace, s); break

                    break

    except Exception as e:

        log_with_time(f"Accès PARTAGE impossible via listdir: {e}")



    if not base_dir or not os.path.isdir(base_dir):

        base_dir = os.path.join(BASE_SHARE, SUBPATH)



    for d in (base_dir, to_long_unc(base_dir)):

        try:

            if not os.path.isdir(d): continue

            files = os.listdir(d)

            qgz = [f for f in files if f.lower().endswith(".qgz")]
            qgz = [f for f in qgz if normalize_name(f).startswith(normalize_name("Contexte éco -"))]
            return [from_long_unc(os.path.join(d, f)) for f in sorted(qgz)]

        except Exception:

            continue

    return []



# =========================

# Fonctions IGN (onglet 2) — identiques au script source

# =========================

def dms_to_dd(text: str) -> float:

    pat = r"(\d{1,3})[°d]\s*(\d{1,2})['m]\s*([\d\.]+)[\"s]?\s*([NSEW])"

    alt = r"(\d{1,3})\s+(\d{1,2})\s+([\d\.]+)\s*([NSEW])"

    m = re.search(pat, text, re.I) or re.search(alt, text, re.I)

    if not m:

        raise ValueError(f"Format DMS invalide : {text}")

    deg, mn, sc, hemi = m.groups()

    dd = float(deg) + float(mn)/60 + float(sc)/3600

    return -dd if hemi.upper() in ("S", "W") else dd



def dd_to_dms(lat: float, lon: float) -> str:

    """Convertit des coordonnées décimales en DMS (degrés, minutes, secondes)."""

    def _convert(value: float, positive: str, negative: str) -> str:

        hemi = positive if value >= 0 else negative

        value = abs(value)

        deg = int(value)

        minutes_full = (value - deg) * 60

        minutes = int(minutes_full)

        seconds = (minutes_full - minutes) * 60

        return f"{deg}°{minutes:02d}'{seconds:04.1f}\"{hemi}"



    return f"{_convert(lat, 'N', 'S')} {_convert(lon, 'E', 'W')}"



def add_hyperlink(paragraph, url: str, text: str, italic: bool = True):

    part = paragraph.part

    r_id = part.relate_to(

        url,

        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",

        is_external=True

    )

    fld_simple = OxmlElement('w:hyperlink')

    fld_simple.set(qn('r:id'), r_id)



    run = OxmlElement('w:r')

    r_pr = OxmlElement('w:rPr')

    if italic:

        i = OxmlElement('w:i')

        r_pr.append(i)

    u = OxmlElement('w:u')

    u.set(qn('w:val'), 'single')

    r_pr.append(u)

    run.append(r_pr)

    run_text = OxmlElement('w:t')

    run_text.text = text

    run.append(run_text)

    fld_simple.append(run)

    paragraph._p.append(fld_simple)



# =========================

# UI — Styles communs

# =========================

class StyleHelper:

    def __init__(self, master, prefs: dict):

        self.master = master

        self.prefs = prefs

        self.style = ttk.Style()

        try: self.style.theme_use("clam")

        except Exception: pass



    def apply(self, theme: str = "light"):

        if theme == "light":

            bg, fg, card_bg, accent, border, subfg = "#F6F7F9", "#111827", "#FFFFFF", "#2563EB", "#E5E7EB", "#6B7280"

            active_accent = "#1D4ED8"

        elif theme == "dark":

            bg, fg, card_bg, accent, border, subfg = "#0F172A", "#E5E7EB", "#111827", "#3B82F6", "#1F2937", "#9CA3AF"

            active_accent = "#2563EB"

        else:  # thème funky

            bg, fg, card_bg, accent, border, subfg = "#1E1E2F", "#FCEFF9", "#27293D", "#FF47A1", "#37394D", "#FFE66D"

            active_accent = "#E0007D"



        self.master.configure(bg=bg)

        s = self.style

        s.configure(".", background=bg, foreground=fg, fieldbackground=card_bg, bordercolor=border)

        # Card-like containers

        s.configure("Card.TFrame", background=card_bg, bordercolor=border, relief="solid", borderwidth=1, padding=8)

        s.configure("Header.TFrame", background=card_bg, bordercolor=border, relief="flat", padding=6)

        s.configure("TLabel", background=bg, foreground=fg)

        s.configure("Card.TLabel", background=card_bg, foreground=fg)

        s.configure("Subtle.TLabel", background=bg, foreground=subfg)

        s.configure("Tooltip.TLabel", background="#111827", foreground="#F9FAFB")

        # Buttons: compact, consistent

        s.configure("Accent.TButton", padding=(12, 6), background=accent, foreground="#FFFFFF")

        s.map("Accent.TButton", background=[("active", active_accent)], foreground=[("active", "#FFFFFF")])

        s.configure("Card.TCheckbutton", background=card_bg, foreground=fg)

        s.configure("Card.TRadiobutton", background=card_bg, foreground=fg)

        s.configure("Card.TEntry", fieldbackground=card_bg)

        s.configure("Status.TLabel", background=card_bg, foreground=subfg)

        s.configure("TProgressbar", troughcolor=border)



# =========================

# (Ancien) onglet Export Cartes

# =========================

class ExportCartesTab(ttk.Frame):

    def __init__(self, parent, style_helper: StyleHelper, prefs: dict):

        super().__init__(parent, padding=12)

        self.parent = parent

        self.prefs = prefs

        self.style_helper = style_helper



        self.font_title = tkfont.Font(family="Segoe UI", size=15, weight="bold")

        self.font_sub   = tkfont.Font(family="Segoe UI", size=10)

        self.font_mono  = tkfont.Font(family="Consolas", size=9)



        self.ze_shp_var   = tk.StringVar(value=from_long_unc(self.prefs.get("ZE_SHP", "")))

        self.ae_shp_var   = tk.StringVar(value=from_long_unc(self.prefs.get("AE_SHP", "")))

        self.cadrage_var  = tk.StringVar(value=self.prefs.get("CADRAGE_MODE", "BOTH"))

        self.overwrite_var= tk.BooleanVar(value=self.prefs.get("OVERWRITE", OVERWRITE_DEFAULT))

        self.dpi_var      = tk.IntVar(value=int(self.prefs.get("DPI", DPI_DEFAULT)))

        self.workers_var  = tk.IntVar(value=int(self.prefs.get("N_WORKERS", N_WORKERS_DEFAULT)))

        self.wait_tiles_var = tk.DoubleVar(value=float(self.prefs.get("WAIT_TILES", WAIT_TILES_DEFAULT)))

        self.commune_var = tk.StringVar(value="")

        self.margin_var   = tk.DoubleVar(value=float(self.prefs.get("MARGIN_FAC", MARGIN_FAC_DEFAULT)))
        self.id_buffer_km_var = tk.DoubleVar(value=float(self.prefs.get("ID_BUFFER_KM", 5.0)))



        self.project_vars: dict[str, tk.IntVar] = {}

        self.all_projects: List[str] = []
        self.filtered_projects: List[str] = []

        self.total_expected = 0

        self.progress_done  = 0
        self.shared_driver = None

        self._build_ui()

        self._populate_projects()

        self._update_counts()

    def _build_ui(self):
        # State
        self.busy = False
        self.wiki_last_url = ""

        # Vars used elsewhere
        self.out_dir_var = tk.StringVar(value=self.prefs.get("OUT_DIR", OUT_IMG))
        # Alias for buffer var used by start_id_thread
        self.buffer_var = tk.DoubleVar(value=float(self.prefs.get("ID_TAMPON_KM", self.id_buffer_km_var.get())))
        # Export type
        self.export_type_var = tk.StringVar(value=self.prefs.get("EXPORT_TYPE", "BOTH"))  # PNG, QGS, BOTH
        # Wiki query manual override
        self.wiki_query_var = tk.StringVar(value=self.prefs.get("WIKI_QUERY", ""))

        # Root layout: left controls, right projects, bottom console
        root = ttk.Frame(self, style="Header.TFrame")
        root.pack(fill="both", expand=True)

        left = ttk.Frame(root, style="Card.TFrame")
        right = ttk.Frame(root, style="Card.TFrame")
        bottom = ttk.Frame(self, style="Card.TFrame")

        left.grid(row=0, column=0, sticky="nsew", padx=6, pady=6)
        right.grid(row=0, column=1, sticky="nsew", padx=6, pady=6)
        bottom.pack(fill="both", expand=True, padx=6, pady=(0,6))

        root.columnconfigure(0, weight=1)
        root.columnconfigure(1, weight=2)
        root.rowconfigure(0, weight=1)

        # --- Left: parameters and actions ---
        ttk.Label(left, text="Contexte éco — Paramètres", font=self.font_title).grid(row=0, column=0, columnspan=3, sticky="w", pady=(0,8))

        # Shapefile selectors
        ttk.Label(left, text="Shapefile Zone d'étude (ZE)").grid(row=1, column=0, sticky="w")
        ze_entry = ttk.Entry(left, textvariable=self.ze_shp_var, width=48, style="Card.TEntry")
        ze_entry.grid(row=1, column=1, sticky="ew", padx=(6,6))
        ttk.Button(left, text="Parcourir…", command=lambda: self._browse_file(self.ze_shp_var)).grid(row=1, column=2)

        ttk.Label(left, text="Shapefile Aire d'étude (AE)").grid(row=2, column=0, sticky="w")
        ae_entry = ttk.Entry(left, textvariable=self.ae_shp_var, width=48, style="Card.TEntry")
        ae_entry.grid(row=2, column=1, sticky="ew", padx=(6,6))
        ttk.Button(left, text="Parcourir…", command=lambda: self._browse_file(self.ae_shp_var)).grid(row=2, column=2)

        # Export options
        opt_frm = ttk.Frame(left)
        opt_frm.grid(row=3, column=0, columnspan=3, sticky="ew", pady=(8,4))
        for i in range(6):
            opt_frm.columnconfigure(i, weight=1)

        ttk.Label(opt_frm, text="DPI").grid(row=0, column=0, sticky="w")
        ttk.Spinbox(opt_frm, from_=72, to=1200, increment=10, textvariable=self.dpi_var, width=6).grid(row=0, column=1, sticky="w")

        ttk.Label(opt_frm, text="Workers").grid(row=0, column=2, sticky="w")
        ttk.Spinbox(opt_frm, from_=1, to=max(1, (os.cpu_count() or 2)), textvariable=self.workers_var, width=4).grid(row=0, column=3, sticky="w")

        ttk.Label(opt_frm, text="Marge (fac.)").grid(row=0, column=4, sticky="w")
        ttk.Spinbox(opt_frm, from_=1.0, to=2.0, increment=0.05, textvariable=self.margin_var, width=6).grid(row=0, column=5, sticky="w")

        ttk.Checkbutton(opt_frm, text="Écraser existants", variable=self.overwrite_var, style="Card.TCheckbutton").grid(row=1, column=0, columnspan=2, sticky="w", pady=(6,0))

        ttk.Label(opt_frm, text="Type d'export").grid(row=1, column=2, sticky="w", pady=(6,0))
        types = ttk.Frame(opt_frm)
        types.grid(row=1, column=3, columnspan=3, sticky="w", pady=(6,0))
        ttk.Radiobutton(types, text="PNG", value="PNG", variable=self.export_type_var, style="Card.TRadiobutton").pack(side="left")
        ttk.Radiobutton(types, text="QGS", value="QGS", variable=self.export_type_var, style="Card.TRadiobutton").pack(side="left", padx=(8,0))
        ttk.Radiobutton(types, text="Les deux", value="BOTH", variable=self.export_type_var, style="Card.TRadiobutton").pack(side="left", padx=(8,0))

        # Output directory
        ttk.Label(left, text="Dossier de sortie").grid(row=4, column=0, sticky="w", pady=(4,0))
        out_entry = ttk.Entry(left, textvariable=self.out_dir_var, width=48, style="Card.TEntry")
        out_entry.grid(row=4, column=1, sticky="ew", padx=(6,6), pady=(4,0))
        ttk.Button(left, text="Choisir…", command=lambda: self._browse_dir(self.out_dir_var)).grid(row=4, column=2, pady=(4,0))

        # Actions: Export, Identification
        act_frm = ttk.Frame(left)
        act_frm.grid(row=5, column=0, columnspan=3, sticky="ew", pady=(10,4))
        act_frm.columnconfigure(0, weight=1)
        act_frm.columnconfigure(1, weight=1)
        self.export_button = ttk.Button(act_frm, text="Exporter cartes", style="Accent.TButton", command=self.start_export_thread)
        self.export_button.grid(row=0, column=0, sticky="ew", padx=(0,6))
        self.id_button = ttk.Button(act_frm, text="ID Contexte éco", command=self.start_id_thread)
        self.id_button.grid(row=0, column=1, sticky="ew")

        # ID buffer
        id_frm = ttk.Frame(left)
        id_frm.grid(row=6, column=0, columnspan=3, sticky="ew", pady=(4,8))
        ttk.Label(id_frm, text="Tampon (km)").pack(side="left")
        ttk.Spinbox(id_frm, from_=0.0, to=20.0, increment=0.5, textvariable=self.buffer_var, width=6).pack(side="left", padx=(6,0))

        # Quick tools: RLT, Maps, BV
        tools_frm = ttk.Frame(left)
        tools_frm.grid(row=7, column=0, columnspan=3, sticky="ew", pady=(2,10))
        self.rlt_button = ttk.Button(tools_frm, text="Remonter le temps (IGN)", command=self.start_rlt_thread)
        self.rlt_button.pack(side="left")
        ttk.Button(tools_frm, text="Google Maps", command=self.open_gmaps).pack(side="left", padx=6)
        self.bassin_button = ttk.Button(tools_frm, text="Bassin versant", command=self.start_bassin_thread)
        self.bassin_button.pack(side="left")

        # Wikipedia + Veg/Sol scraping controls
        wiki_frm = ttk.LabelFrame(left, text="Scraping Wikipedia / Veg&Sol")
        wiki_frm.grid(row=8, column=0, columnspan=3, sticky="ew")
        ttk.Label(wiki_frm, text="Requête manuelle (facultatif)").grid(row=0, column=0, sticky="w", padx=6, pady=4)
        ttk.Entry(wiki_frm, textvariable=self.wiki_query_var, width=40).grid(row=0, column=1, sticky="ew", padx=6, pady=4)
        self.wiki_button = ttk.Button(wiki_frm, text="Lancer scraping", command=lambda: threading.Thread(target=self._run_all_scrapers, daemon=True).start())
        self.wiki_button.grid(row=1, column=0, sticky="w", padx=6, pady=(0,6))
        self.wiki_open_button = ttk.Button(wiki_frm, text="Ouvrir la page Wikipédia", command=self._open_wiki_url, state="disabled")
        self.wiki_open_button.grid(row=1, column=1, sticky="w", padx=6, pady=(0,6))
        try:
            ToolTip(self.wiki_button, "Scrape Wikipedia et Veg&Sol en utilisant la ZE/AE")
        except Exception:
            pass

        for c in range(3):
            left.columnconfigure(c, weight=1)

        # --- Right: project list with filter ---
        ttk.Label(right, text="Projets QGIS (Contexte éco)", font=self.font_title).pack(anchor="w", pady=(0,6))
        filt_frm = ttk.Frame(right)
        filt_frm.pack(fill="x", pady=(0,4))
        ttk.Label(filt_frm, text="Filtrer").pack(side="left")
        self.filter_var = tk.StringVar(value="")
        ent = ttk.Entry(filt_frm, textvariable=self.filter_var, width=30)
        ent.pack(side="left", padx=(6,6))
        ttk.Button(filt_frm, text="Appliquer", command=self._apply_filter).pack(side="left")
        ttk.Button(filt_frm, text="Rafraîchir", command=self._populate_projects).pack(side="left", padx=(6,0))
        ttk.Button(filt_frm, text="Tout", command=lambda: self._select_all(True)).pack(side="left", padx=(12,0))
        ttk.Button(filt_frm, text="Aucun", command=lambda: self._select_all(False)).pack(side="left", padx=(6,0))

        # Scrollable area
        canvas = tk.Canvas(right, highlightthickness=0, bd=0)
        vsb = ttk.Scrollbar(right, orient="vertical", command=canvas.yview)
        container = ttk.Frame(canvas)
        container.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        self.scrollable_frame = ttk.Frame(container)
        self.scrollable_frame.pack(fill="both", expand=True)
        canvas_frame = canvas.create_window((0, 0), window=container, anchor="nw")
        canvas.configure(yscrollcommand=vsb.set)
        canvas.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")

        def _resize_container(event):
            canvas.itemconfigure(canvas_frame, width=event.width)
        canvas.bind("<Configure>", _resize_container)

        # --- Bottom: results + console + progress ---
        top_bottom = ttk.Frame(bottom)
        top_bottom.pack(fill="both", expand=True)

        # Results tree (2 columns)
        res_frm = ttk.Frame(top_bottom)
        res_frm.pack(side="left", fill="both", expand=False, padx=(0,6))
        ttk.Label(res_frm, text="Résultats (Wikipédia)").pack(anchor="w")
        self.results_tree = ttk.Treeview(res_frm, columns=("label", "text"), show="headings", height=6)
        self.results_tree.heading("label", text="Rubrique")
        self.results_tree.heading("text", text="Texte")
        self.results_tree.column("label", width=160, anchor="w")
        self.results_tree.column("text", width=480, anchor="w")
        self.results_tree.pack(fill="both", expand=True)
        # Initialize rows
        for iid, lbl in (("climat", "Climat"), ("occupation_sols", "Occupation des sols"), ("altitude", "Altitude"), ("vegetation", "Végétation"), ("sols", "Sols")):
            try:
                self.results_tree.insert("", "end", iid=iid, values=(lbl, ""))
            except Exception:
                pass

        # Console
        cons_frm = ttk.Frame(top_bottom)
        cons_frm.pack(side="left", fill="both", expand=True)
        ttk.Label(cons_frm, text="Console").pack(anchor="w")
        self.console_text = tk.Text(cons_frm, height=10, wrap="word", font=self.font_mono, state="disabled")
        self.console_text.pack(fill="both", expand=True)
        self.stdout_redirect = TextRedirector(self.console_text)

        # Progress + status
        status_frm = ttk.Frame(bottom)
        status_frm.pack(fill="x", pady=(6,0))
        self.progress = ttk.Progressbar(status_frm, mode="determinate")
        self.progress.pack(fill="x")
        self.status_label = ttk.Label(status_frm, text="Prêt", style="Status.TLabel")
        self.status_label.pack(anchor="w")

    # Small helpers for file/directory browsing
    def _browse_file(self, var: tk.StringVar):
        try:
            initial = var.get() or DEFAULT_SHAPE_DIR
        except Exception:
            initial = DEFAULT_SHAPE_DIR
        path = filedialog.askopenfilename(title="Sélectionner un shapefile", initialdir=initial, filetypes=[["Shapefile", ".shp .SHP"], ["Tous", "*.*"]])
        if path:
            var.set(path)

    def _browse_dir(self, var: tk.StringVar):
        try:
            initial = var.get() or OUT_IMG
        except Exception:
            initial = OUT_IMG
        path = filedialog.askdirectory(title="Choisir le dossier de sortie", initialdir=initial)
        if path:
            var.set(path)

    def _open_wiki_url(self):
        url = (self.wiki_last_url or "").strip()
        if not url:
            messagebox.showinfo("Wikipedia", "Aucune URL disponible. Lancez d'abord le scraping.")
            return

        print(f"[Wiki] Ouverture dans le navigateur : {url}", file=self.stdout_redirect)
        driver = self._get_or_create_driver()
        if not driver:
            return

        try:
            # Open in a new tab and switch to it
            driver.execute_script(f"window.open('{url}', '_blank');")
            driver.switch_to.window(driver.window_handles[-1])
        except Exception as e:
            messagebox.showerror("Wikipedia", f"Impossible d'ouvrir l'URL : {e}")

    def _get_or_create_driver(self):
        if self.shared_driver is None:
            try:
                # Setup Chrome options
                options = webdriver.ChromeOptions()
                options.add_experimental_option("detach", True)
                options.add_argument("--start-maximized")
                options.add_experimental_option('excludeSwitches', ['enable-logging'])

                if os.getenv('APP_HEADLESS') == '1':
                    options.add_argument('--headless')
                    options.add_argument('--no-sandbox')
                    options.add_argument('--disable-dev-shm-usage')

                self.shared_driver = webdriver.Chrome(options=options)
                print("[WebDriver] Shared browser instance created.", file=self.stdout_redirect)
            except Exception as e:
                messagebox.showerror("Erreur WebDriver", f"Impossible de démarrer le navigateur Chrome partagé : {e}")
                self.shared_driver = None
                return None
        return self.shared_driver

    def _cleanup_driver(self):
        if self.shared_driver:
            try:
                self.shared_driver.quit()
                print("[WebDriver] Shared browser instance closed.", file=self.stdout_redirect)
            except Exception as e:
                print(f"[WebDriver] Error closing shared browser: {e}", file=self.stdout_redirect)
            finally:
                self.shared_driver = None

    def _run_all_scrapers(self):
        os.makedirs("output", exist_ok=True)
        driver = self._get_or_create_driver()
        if not driver:
            print("Cancel scraping due to browser start failure", file=self.stdout_redirect)
            return
        try:
            # Wikipedia scraping in first tab
            self._run_wiki(driver)
            
            # VegSol scraping in new tab
            driver.execute_script("window.open('about:blank', '_blank');")
            driver.switch_to.window(driver.window_handles[-1])
            self._run_vegsol(driver)
            
        except Exception as e:
            print(f"Scraping error: {e}", file=self.stdout_redirect)
        finally:
            print("Scraping finished. Browser remains open for inspection.", file=self.stdout_redirect)
            def _reenable_ui():
                try:
                    self.wiki_button.config(state="normal")
                    if self.wiki_last_url:
                        self.wiki_open_button.config(state="normal")
                except Exception:
                    pass
            self.after(0, _reenable_ui)

    def _run_wiki(self, driver):
        # This method now receives the shared driver
        # and operates in the currently active tab.
        try:
            print("[Wiki] Lancement du scraping Wikipedia", file=self.stdout_redirect)

            ze_path = self.ze_shp_var.get()

            gdf = gpd.read_file(ze_path)

            if gdf.crs is None:

                raise ValueError("CRS non défini")

            gdf = gdf.to_crs("EPSG:4326")

            centroid = gdf.geometry.unary_union.centroid

            lat, lon = centroid.y, centroid.x

            manual = (self.wiki_query_var.get() or "").strip()

            if manual:

                query = manual

                try:

                    self.prefs["WIKI_QUERY"] = manual; save_prefs(self.prefs)

                except Exception:

                    pass

            else:

                commune, dep = self._detect_commune(lat, lon)

                query = f"{commune} {dep}".strip()

            print(f"[Wiki] Requête : {query}", file=self.stdout_redirect)

            data = get_wikipedia_extracts(query)

            # Mettre à jour le tableau Wikipedia (dès que les données sont disponibles)

            self._update_wiki_table(data)

            if "error" in data:

                print(f"[Wiki] {data['error']}", file=self.stdout_redirect)

            else:

                print(f"[Wiki] Page Wikipédia : {data['url']}", file=self.stdout_redirect)

                print("[Wiki] CLIMAT :", file=self.stdout_redirect)

                if data['climat_p1'] != 'Non trouvé':

                    print(data['climat_p1'], file=self.stdout_redirect)

                if data['climat_p2'] != 'Non trouvé':

                    print(data['climat_p2'], file=self.stdout_redirect)

                print("[Wiki] OCCUPATION DES SOLS :", file=self.stdout_redirect)

                if data['occupation_p1'] != 'Non trouvé':

                    print(data['occupation_p1'], file=self.stdout_redirect)

        except Exception as e:

            print(f"[Wiki] Erreur : {e}", file=self.stdout_redirect)




    def _update_wiki_table(self, data: dict) -> None:

        try:
            # Collect values from either legacy or new keys
            clim_txt = data.get('climat') or data.get('climat_p1') or ''
            occ_txt  = data.get('occup_sols') or data.get('occupation_p1') or ''
            url_txt  = data.get('url', '')

            payload = {
                'climat': clim_txt if clim_txt and not str(clim_txt).lower().startswith('non trouv') else 'Non trouvé',
                'occupation_sols': occ_txt if occ_txt and not str(occ_txt).lower().startswith('non trouv') else 'Non trouvé',
            }

            # Update the consolidated results tree on the UI thread
            self.after(0, self._update_results_tree, payload)

            # Store URL and enable the open button if present
            def _upd_url():
                try:
                    self.wiki_last_url = url_txt or ""
                    btn = getattr(self, 'wiki_open_button', None) or getattr(self, 'open_scraping_button', None)
                    if btn:
                        btn.config(state=("normal" if self.wiki_last_url else "disabled"))
                except Exception:
                    pass
            self.after(0, _upd_url)

        except Exception:

            pass



    def _update_results_tree(self, data: dict) -> None:
        """Update rows in self.results_tree. Keys are expected to be iids:
        'climat', 'occupation_sols', 'altitude', 'vegetation', 'sols'.
        """
        try:
            if not hasattr(self, 'results_tree'):
                return
            labels = {
                'climat': 'Climat',
                'occupation_sols': 'Occupation sols',
                'altitude': 'Altitude',
                'vegetation': 'Végétation',
                'sols': 'Sols',
            }
            for iid, text in (data or {}).items():
                if iid not in labels:
                    continue
                try:
                    # Preserve first column label
                    cur = self.results_tree.item(iid, 'values')
                    label = cur[0] if (isinstance(cur, (list, tuple)) and len(cur) >= 1) else labels[iid]
                    self.results_tree.item(iid, values=(label, text))
                except Exception:
                    # If row doesn't exist yet, insert it
                    self.results_tree.insert("", "end", values=(labels[iid], text), iid=iid)
        except Exception:
            pass






    def _run_vegsol(self, driver):
        # This method now receives the shared driver
        # and operates in a new tab as orchestrated by _run_all_scrapers.

        try:

            print("[Cartes] Lancement du scraping des cartes", file=self.stdout_redirect)

            ze_path = self.ze_shp_var.get()

            gdf = gpd.read_file(ze_path)

            if gdf.crs is None:

                raise ValueError("CRS non défini")

            gdf = gdf.to_crs("EPSG:4326")

            centroid = gdf.geometry.unary_union.centroid

            lat, lon = centroid.y, centroid.x

            coords_dms = dd_to_dms(lat, lon)

            options = webdriver.ChromeOptions()

            options.add_experimental_option("excludeSwitches", ["enable-logging"])

            options.add_argument("--log-level=3")

            options.add_argument("--disable-extensions")

            options.add_argument("--disable-gpu")

            options.add_argument("--no-sandbox")

            options.add_argument("--disable-dev-shm-usage")

            # Respect APP_HEADLESS env var (default: visible)

            try:

                if os.environ.get("APP_HEADLESS", "0").lower() in ("1", "true", "yes"):

                    options.add_argument("--headless=new")

            except Exception:

                try:

                    if os.environ.get("APP_HEADLESS", "0").lower() in ("1", "true", "yes"):

                        options.add_argument("--headless")

                except Exception:

                    pass

            # Driver local si présent

            local_driver = os.path.join(REPO_ROOT if 'REPO_ROOT' in globals() else os.path.abspath(os.path.join(os.path.dirname(__file__), '..')), 'tools', 'chromedriver.exe')

            if os.path.isfile(local_driver):

                self.vegsol_driver = webdriver.Chrome(service=Service(local_driver), options=options)

            else:

                self.vegsol_driver = webdriver.Chrome(options=options)

            self.vegsol_driver.maximize_window()



            # Nouveau flux automatisé (import shapefile + scraping pop-up)

            try:

                wait = WebDriverWait(self.vegsol_driver, 10)

                # 1) Ouvrir l'URL

                self.vegsol_driver.get("https://floreapp.netlify.app/biblio-patri.html")

                time.sleep(0.75)



                # 3) Cliquer sur Importer shapefile

                try:

                    btn_upload = wait.until(EC.element_to_be_clickable((By.ID, "upload-shapefile-btn")))

                    btn_upload.click()

                except Exception:

                    pass

                time.sleep(0.75)



                # 5) Cliquer sur Zone d’étude

                try:

                    btn_zone = wait.until(EC.element_to_be_clickable((By.ID, "import-zone-btn")))

                    btn_zone.click()

                except Exception:

                    pass



                # Préparer la liste des fichiers du shapefile

                def _from_long_unc(p: str) -> str:

                    p = p or ""

                    if p.startswith("\\\\?\\UNC"):
                        return "\\\\" + p[8:]

                    if p.startswith("\\\\?\\"):
                        return p[4:]

                    return p



                ze_shp = (self.ze_shp_var.get() or "").strip()

                base_no_ext, _ = os.path.splitext(_from_long_unc(ze_shp))

                exts = [".cpg", ".dbf", ".prj", ".qmd", ".shp", ".shx"]

                files = [base_no_ext + e for e in exts if os.path.isfile(base_no_ext + e)]

                if not files:

                    raise ValueError("Fichiers du shapefile introuvables pour l'import")



                # Envoyer les fichiers à l'input[type=file]

                inputs = self.vegsol_driver.find_elements(By.CSS_SELECTOR, "input[type='file']")

                target_input = inputs[-1] if inputs else None

                if not target_input:

                    raise RuntimeError("Champ d'import fichier introuvable")

                try:

                    target_input.send_keys("\n".join(files))

                except Exception as e:

                    print(f"[Cartes] Envoi fichiers échoué: {e}", file=self.stdout_redirect)



                time.sleep(0.75)



                # 8) Clic droit au centre de la carte

                map_el = wait.until(EC.visibility_of_element_located((By.ID, "map")))

                ActionChains(self.vegsol_driver).move_to_element(map_el).context_click(map_el).perform()

                time.sleep(0.5)



                # 10) Cliquer sur 'Ressources'
                print("[Cartes] Recherche du bouton 'Ressources'...", file=self.stdout_redirect)
                
                try:
                    # Essayer plusieurs sélecteurs pour le bouton Ressources
                    ressources_selectors = [
                        "//button[contains(.,'Ressources')]",
                        "//button[contains(text(),'Ressources')]",
                        "//button[@class='action-button' and contains(.,'Ressources')]",
                        ".action-button"
                    ]
                    
                    btn_res = None
                    for selector in ressources_selectors:
                        try:
                            if selector.startswith("//"):
                                btn_res = wait.until(EC.element_to_be_clickable((By.XPATH, selector)))
                            else:
                                btn_res = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, selector)))
                            print(f"[Cartes] Bouton trouvé avec sélecteur: {selector}", file=self.stdout_redirect)
                            break
                        except Exception:
                            continue
                    
                    if btn_res:
                        print("[Cartes] Clic sur le bouton Ressources...", file=self.stdout_redirect)
                        btn_res.click()
                    else:
                        print("[Cartes] Bouton Ressources non trouvé - recherche de tous les boutons...", file=self.stdout_redirect)
                        # Debug: lister tous les boutons disponibles
                        buttons = self.vegsol_driver.find_elements(By.TAG_NAME, "button")
                        for i, btn in enumerate(buttons):
                            try:
                                text = btn.text.strip()
                                if text:
                                    print(f"[Cartes] Bouton {i}: '{text}'", file=self.stdout_redirect)
                            except Exception:
                                pass

                except Exception as btn_err:
                    print(f"[Cartes] Erreur bouton Ressources: {btn_err}", file=self.stdout_redirect)

                time.sleep(2.0)  # Attendre plus longtemps pour la popup



                # 12) Scraper les éléments de la pop-up avec attente améliorée
                
                print("[Cartes] Attente de la popup avec les données...", file=self.stdout_redirect)
                
                # Attendre plus longtemps pour que la popup se charge complètement
                time.sleep(2.0)
                
                # Essayer plusieurs fois de trouver les éléments
                alt = veg = soil = "Non trouvé"
                
                for attempt in range(3):
                    try:
                        print(f"[Cartes] Tentative {attempt + 1} de scraping...", file=self.stdout_redirect)
                        
                        html = self.vegsol_driver.page_source
                        soup = BeautifulSoup(html, "lxml")
                        
                        # Debug: afficher les sélecteurs trouvés
                        altitude_els = soup.select("div.altitude-info")
                        veg_els = soup.select("div.tooltip-pill.vegetation-pill")
                        soil_els = soup.select("div.tooltip-pill.soil-pill")
                        
                        print(f"[Cartes] Éléments trouvés - Altitude: {len(altitude_els)}, Végétation: {len(veg_els)}, Sols: {len(soil_els)}", file=self.stdout_redirect)
                        
                        # Debug: afficher le contenu HTML des éléments trouvés
                        if altitude_els:
                            print(f"[Cartes] Contenu altitude HTML: {str(altitude_els[0])[:200]}", file=self.stdout_redirect)
                        if veg_els:
                            for i, el in enumerate(veg_els):
                                print(f"[Cartes] Contenu végétation {i} HTML: {str(el)[:200]}", file=self.stdout_redirect)
                        if soil_els:
                            for i, el in enumerate(soil_els):
                                print(f"[Cartes] Contenu sols {i} HTML: {str(el)[:200]}", file=self.stdout_redirect)
                        
                        def _txt(sel):
                            el = soup.select_one(sel)
                            if el:
                                text = el.get_text(" ", strip=True)
                                print(f"[Cartes] Texte extrait pour {sel}: '{text}'", file=self.stdout_redirect)
                                return text if text.strip() else "Non trouvé"
                            return "Non trouvé"
                        
                        alt_new = _txt("div.altitude-info")
                        veg_new = _txt("div.tooltip-pill.vegetation-pill")
                        soil_new = _txt("div.tooltip-pill.soil-pill")
                        
                        # Si on trouve au moins un élément valide, on garde les résultats
                        if alt_new != "Non trouvé":
                            alt = alt_new
                        if veg_new != "Non trouvé":
                            veg = veg_new
                        if soil_new != "Non trouvé":
                            soil = soil_new
                            
                        # Si on a trouvé tous les éléments, on peut arrêter
                        if alt != "Non trouvé" and veg != "Non trouvé" and soil != "Non trouvé":
                            break
                            
                        # Sinon, attendre un peu plus
                        if attempt < 2:
                            time.sleep(1.5)
                            
                    except Exception as scrape_err:
                        print(f"[Cartes] Erreur scraping tentative {attempt + 1}: {scrape_err}", file=self.stdout_redirect)
                        if attempt < 2:
                            time.sleep(1.5)

                # Update the consolidated results table
                payload = {
                    'altitude': alt,
                    'vegetation': veg, 
                    'sols': soil
                }
                self.after(0, self._update_results_tree, payload)
                
                print(f"[Cartes] ALTITUDE : {alt}", file=self.stdout_redirect)
                print(f"[Cartes] VÉGÉTATION : {veg}", file=self.stdout_redirect)
                print(f"[Cartes] SOLS : {soil}", file=self.stdout_redirect)

                return

            except Exception as flow_err:

                print(f"[Cartes] Flux import/scraping échoué: {flow_err}", file=self.stdout_redirect)




        except Exception as e:

            print(f"[Cartes] Erreur : {e}", file=self.stdout_redirect)

        finally:
            try:
                if hasattr(self, 'vegsol_driver'):
                    self.vegsol_driver.quit()
            except Exception:
                pass

    # --- Boutons ajoutés ---


    def open_biodiv_dialog(self) -> None:
        try:
            if hasattr(self, 'biodiv_win') and self.biodiv_win and tk.Toplevel.winfo_exists(self.biodiv_win):
                self.biodiv_win.lift()
                return
        except Exception:
            pass

        win = tk.Toplevel(self)
        win.title("Photo Biodiv'AURA")
        win.geometry("640x520")
        self.biodiv_win = win

        frm = ttk.Frame(win, padding=12)
        frm.pack(fill=tk.BOTH, expand=True)

        ttk.Label(frm, text="Collez 1 espèce par ligne (20 max)", style="Card.TLabel").pack(anchor='w')

        # Scrollable text area with mouse wheel support
        area = ttk.Frame(frm)
        area.pack(fill=tk.BOTH, expand=True, pady=(6,6))
        txt = tk.Text(area, height=20, wrap=tk.NONE)
        yscroll = ttk.Scrollbar(area, orient="vertical", command=txt.yview)
        txt.configure(yscrollcommand=yscroll.set)
        area.rowconfigure(0, weight=1)
        area.columnconfigure(0, weight=1)
        txt.grid(row=0, column=0, sticky="nsew")
        yscroll.grid(row=0, column=1, sticky="ns")
        self.biodiv_text = txt
        def _mw_text(e):
            try:
                delta = -1 * (e.delta // 120)
            except Exception:
                delta = -1 if getattr(e, 'num', 0) == 4 else (1 if getattr(e, 'num', 0) == 5 else 0)
            if delta:
                txt.yview_scroll(delta, "units")
            return "break"
        txt.bind("<MouseWheel>", _mw_text)
        txt.bind("<Button-4>", _mw_text)
        txt.bind("<Button-5>", _mw_text)

        btns = ttk.Frame(frm)
        btns.pack(fill=tk.X)

        self.biodiv_launch_btn = ttk.Button(btns, text="Lancer le scraping", style="Accent.TButton",
                                            command=self.start_biodiv_thread)
        self.biodiv_launch_btn.pack(side=tk.LEFT)
        try:
            self.biodiv_launch_btn.config(state='normal')
        except Exception:
            pass

        # Permettre Entrée pour lancer rapidement
        try:
            win.bind('<Return>', lambda e: self.start_biodiv_thread())
        except Exception:
            pass

        ttk.Button(btns, text="Fermer", command=win.destroy).pack(side=tk.RIGHT)

    def start_biodiv_thread(self) -> None:
        try:
            print("[Biodiv] Bouton 'Lancer le scraping' cliqué", file=self.stdout_redirect)
        except Exception:
            pass
        try:
            raw = self.biodiv_text.get('1.0', tk.END)
        except Exception:
            raw = ''

        species = [s.strip() for s in (raw.splitlines() if raw else [])]
        species = [s for s in species if s]
        if not species:
            messagebox.showerror("Biodiv'AURA", "Veuillez saisir au moins une espèce.")
            return

        if len(species) > 20:
            species = species[:20]

        print(f"[Biodiv] Lancement pour {len(species)} espèce(s)", file=self.stdout_redirect)

        try:
            if hasattr(self, 'biodiv_launch_btn'):
                self.biodiv_launch_btn.config(state='disabled')
        except Exception:
            pass

        t = threading.Thread(target=self._run_biodiv, args=(species,))
        t.daemon = True
        t.start()

    def _run_biodiv(self, species_list: list[str]) -> None:
        try:
            os.makedirs(OUT_IMG, exist_ok=True)
            out_dir = os.path.join(OUT_IMG, "Photo BiodivAURA")
            os.makedirs(out_dir, exist_ok=True)

            # Document unique pour toutes les especes
            doc = Document()
            try:
                style_normal = doc.styles['Normal']
                style_normal.font.name = 'Calibri'
                style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            except Exception:
                pass

            # Préparer Selenium (toujours visible pour cette fonctionnalité)
            driver = None
            wait = None
            try:
                options = webdriver.ChromeOptions()
                options.add_experimental_option("excludeSwitches", ["enable-logging"]) 
                options.add_argument("--log-level=3")
                options.add_argument("--disable-extensions")
                options.add_argument("--disable-gpu")
                options.add_argument("--no-sandbox")
                options.add_argument("--disable-dev-shm-usage")
                try:
                    # Ne pas attendre le chargement complet des pages
                    options.page_load_strategy = 'none'
                except Exception:
                    pass
                # IMPORTANT: ne pas forcer headless ici afin que l'utilisateur voie le navigateur
                local_driver = os.path.join(REPO_ROOT if 'REPO_ROOT' in globals() else os.path.abspath(os.path.join(os.path.dirname(__file__), '..')), 'tools', 'chromedriver.exe')
                if os.path.isfile(local_driver):
                    driver = webdriver.Chrome(service=Service(local_driver), options=options)
                else:
                    driver = webdriver.Chrome(options=options)
                try:
                    driver.maximize_window()
                except Exception:
                    pass
                # Attentes Selenium courtes mais un peu plus tolérantes (4s)
                wait = WebDriverWait(driver, 4)
            except Exception as se_init:
                print(f"[Biodiv] Selenium init KO: {se_init}", file=self.stdout_redirect)
            print(f"[Biodiv] Scraping de {len(species_list)} espece(s)...", file=self.stdout_redirect)

            # Collecter toutes les données d'espèces avec leurs images
            species_data = []
            
            # Charger le mapping TAXREF (nom latin -> CD_NOM) pour construire directement les URLs
            taxref_map = {}
            try:
                repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir))
                taxref_path = os.path.join(repo_root, "Bases de données", "taxref.json")
                with open(taxref_path, "r", encoding="utf-8") as f:
                    taxref_raw = json.load(f)
                # Normaliser en minuscule; ignorer l'en-tête "nom latin": "CD_NOM"
                for k, v in taxref_raw.items():
                    if k.strip().lower() == "nom latin":
                        continue
                    taxref_map[k.strip().lower()] = str(v).strip()
                print(f"[Biodiv] TAXREF charge: {len(taxref_map)} entrées", file=self.stdout_redirect)
            except Exception as te:
                print(f"[Biodiv] Impossible de charger TAXREF: {te}", file=self.stdout_redirect)
            
            # On n'ouvre pas la page d'accueil ici. On ira directement sur l'URL espèce quand on l'a.

            for idx, sp in enumerate(species_list, start=1):
                try:
                    print(f"[Biodiv] ({idx}/{len(species_list)}) {sp}", file=self.stdout_redirect)
                    
                    # Essayer d'abord de construire l'URL via TAXREF (nom latin -> CD_NOM)
                    url_sp = None
                    cd_nom = None
                    tmp_path = None
                    try:
                        key = sp.strip().lower()
                        cd_nom = taxref_map.get(key)
                    except Exception:
                        cd_nom = None

                    # Si l'utilisateur a collé directement un CD_NOM (ligne numérique)
                    if not cd_nom:
                        stripped = sp.strip().replace(" ", "")
                        if stripped.isdigit():
                            cd_nom = stripped

                    if cd_nom:
                        # Construire l'URL espèce directement
                        url_sp = f"https://atlas.biodiversite-auvergne-rhone-alpes.fr/espece/{cd_nom}"
                        # Utiliser requests en priorité (plus rapide, pas d'attente Selenium)
                        img_url = None
                        tmp_path = None
                        try:
                            resp = requests.get(
                                url_sp,
                                timeout=8,
                                headers={
                                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"
                                },
                            )
                            if resp.ok:
                                soup = BeautifulSoup(resp.text, "html.parser")
                                m = soup.select_one("meta[property='og:image']")
                                if m and m.get("content"):
                                    img_url = m.get("content")
                        except Exception as e_req:
                            print(f"[Biodiv] Echec requete: {e_req}", file=self.stdout_redirect)

                        # Fallback Selenium si pas d'URL image trouvée via requests
                        if not img_url and driver is not None:
                            try:
                                driver.get(url_sp)
                                # Attendre un minimum que le DOM soit prêt
                                try:
                                    WebDriverWait(driver, 4).until(lambda d: d.execute_script("return document.readyState") in ("interactive", "complete"))
                                except Exception:
                                    pass
                                # Extraire og:image via JS avec échappement correct
                                og = driver.execute_script('var m=document.querySelector("meta[property=\'og:image\']"); return m?m.content:null;')
                                if og and isinstance(og, str) and og.startswith("http"):
                                    img_url = og
                            except Exception as e_se:
                                print(f"[Biodiv] Selenium fallback KO: {e_se}", file=self.stdout_redirect)

                        # Télécharger l'image si une URL a été trouvée
                        if img_url:
                            try:
                                img_resp = requests.get(
                                    img_url,
                                    stream=True,
                                    timeout=15,
                                    headers={
                                        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"
                                    },
                                )
                                if img_resp.ok:
                                    safe_name = re.sub(r"[^\w\-]+", "_", sp.strip())[:80]
                                    ext = os.path.splitext((img_url.split("?")[0]).strip())[1] or ".jpg"
                                    if len(ext) > 6:
                                        ext = ".jpg"
                                    tmp_path = os.path.join(out_dir, f"{safe_name}{ext}")
                                    with open(tmp_path, "wb") as f:
                                        for chunk in img_resp.iter_content(chunk_size=8192):
                                            if chunk:
                                                f.write(chunk)
                                    print(f"[Biodiv] Image ok: {tmp_path}", file=self.stdout_redirect)
                                else:
                                    print(f"[Biodiv] HTTP {img_resp.status_code} pour {img_url}", file=self.stdout_redirect)
                            except Exception as de_dl:
                                print(f"[Biodiv] Echec dl image {img_url}: {de_dl}", file=self.stdout_redirect)
                        else:
                            print(f"[Biodiv] Aucune image pour {sp} ({url_sp or 'URL inconnue'})", file=self.stdout_redirect)
                    

                    # URL de la page espece (deja calculee si TAXREF)
                    if not url_sp:
                        try:
                            url_sp = driver.current_url
                        except Exception:
                            url_sp = None

                    species_data.append({
                        'name': sp,
                        'image_path': tmp_path,
                        'url': url_sp
                    })

                except Exception as sp_err:
                    print(f"[Biodiv] Erreur espece {sp}: {sp_err}", file=self.stdout_redirect)
                    species_data.append({'name': sp, 'image_path': None, 'url': None})

            if driver is not None:
                try:
                    driver.quit()
                except Exception:
                    pass

            # Créer le tableau avec les images (2x3 format)
            self._create_species_table(doc, species_data)

            ts = datetime.datetime.now().strftime('%Y%m%d-%H%M')
            doc_path = os.path.join(out_dir, f"Photos_BiodivAURA_{ts}.docx")
            try:
                doc.save(doc_path)
                print(f"[Biodiv] Document genere: {doc_path}", file=self.stdout_redirect)
            except Exception as se:
                print(f"[Biodiv] Sauvegarde echouee: {se}", file=self.stdout_redirect)

            self._set_status(f"Termine - {doc_path}")

        except Exception as e:
            print(f"[Biodiv] Erreur: {e}", file=self.stdout_redirect)
        finally:
            try:
                if hasattr(self, 'biodiv_launch_btn'):
                    self.after(0, lambda: self.biodiv_launch_btn.config(state='normal'))
            except Exception:
                pass

    def _create_species_table(self, doc, species_data):
        """Crée un tableau 2x3 avec les images d'espèces comme dans le screenshot"""
        if not species_data:
            return

        # Traiter les espèces par groupes de 6 (2 lignes x 3 colonnes)
        for page_start in range(0, len(species_data), 6):
            page_species = species_data[page_start:page_start + 6]
            
            # Calculer le nombre de lignes nécessaires
            rows_needed = (len(page_species) + 2) // 3  # +2 pour arrondir vers le haut
            
            # Créer le tableau
            table = doc.add_table(rows=rows_needed, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Configurer le style du tableau
            table.style = 'Table Grid'
            
            # Remplir le tableau
            for i, species in enumerate(page_species):
                row_idx = i // 3
                col_idx = i % 3
                cell = table.cell(row_idx, col_idx)
                
                # Vider la cellule
                cell.text = ''
                
                # Ajouter l'image
                if species['image_path'] and os.path.exists(species['image_path']):
                    try:
                        # Créer un paragraphe pour l'image
                        img_paragraph = cell.paragraphs[0]
                        img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = img_paragraph.runs[0] if img_paragraph.runs else img_paragraph.add_run()
                        
                        # Ajuster la taille de l'image pour le tableau (plus petite)
                        img_width = Cm(4.5)  # Largeur réduite pour s'adapter au tableau
                        run.add_picture(species['image_path'], width=img_width)
                        
                    except Exception as e:
                        print(f"[Biodiv] Erreur ajout image {species['name']}: {e}", file=self.stdout_redirect)
                        cell.paragraphs[0].add_run("[Image non disponible]")
                else:
                    cell.paragraphs[0].add_run("[Image non disponible]")
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Ajouter le nom de l'espèce en dessous
                name_paragraph = cell.add_paragraph()
                name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                name_run = name_paragraph.add_run(species['name'])
                name_run.bold = True
                
                # Ajouter le lien si disponible
                if species['url']:
                    try:
                        link_paragraph = cell.add_paragraph()
                        link_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        add_hyperlink(link_paragraph, species['url'], "Voir sur Biodiv'AURA", italic=True)
                    except Exception:
                        pass
            
            # Ajuster la largeur des cellules
            for row in table.rows:
                for cell in row.cells:
                    # Définir la largeur des cellules
                    cell.width = Cm(5.5)
            
            # Ajouter un saut de page si ce n'est pas la dernière page
            if page_start + 6 < len(species_data):
                doc.add_page_break()



    def open_google_maps(self):
        """Ouvre Google Maps avec le centroïde de la zone d'étude"""
        ze_path = self.ze_shp_var.get()
        if not ze_path or not os.path.isfile(ze_path):
            messagebox.showerror("Erreur", "Veuillez d'abord sélectionner un shapefile pour la Zone d'étude.")
            return
        
        try:
            import geopandas as gpd
            gdf = gpd.read_file(ze_path)
            if gdf.crs is None:
                messagebox.showwarning("Avertissement", "Le shapefile n'a pas de CRS défini.")
            
            # Reproject to WGS84 for lat/lon
            gdf_wgs84 = gdf.to_crs("EPSG:4326")
            centroid = gdf_wgs84.geometry.unary_union.centroid
            lat, lon = centroid.y, centroid.x
            
            url = f"https://www.google.com/maps/@{lat},{lon},15z"
            import webbrowser
            webbrowser.open(url)
            print(f"[Maps] Ouverture Google Maps: {lat:.6f}, {lon:.6f}", file=self.stdout_redirect)
            
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de l'ouverture de Google Maps : {e}")

    def open_remonter_temps(self):
        """Ouvre Remonter le temps avec le centroïde de la zone d'étude"""
        ze_path = self.ze_shp_var.get()
        if not ze_path or not os.path.isfile(ze_path):
            messagebox.showerror("Erreur", "Veuillez d'abord sélectionner un shapefile pour la Zone d'étude.")
            return
        
        try:
            import geopandas as gpd
            gdf = gpd.read_file(ze_path)
            if gdf.crs is None:
                messagebox.showwarning("Avertissement", "Le shapefile n'a pas de CRS défini.")
            
            # Reproject to WGS84 for lat/lon
            gdf_wgs84 = gdf.to_crs("EPSG:4326")
            centroid = gdf_wgs84.geometry.unary_union.centroid
            lat, lon = centroid.y, centroid.x
            
            url = f"https://remonterletemps.ign.fr/comparer/basic?x={lon}&y={lat}&z=15&layer1=GEOGRAPHICALGRIDSYSTEMS.PLANIGNV2&layer2=ORTHOIMAGERY.ORTHOPHOTOS&mode=vSlider"
            import webbrowser
            webbrowser.open(url)
            print(f"[Temps] Ouverture Remonter le temps: {lat:.6f}, {lon:.6f}", file=self.stdout_redirect)
            
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de l'ouverture de Remonter le temps : {e}")

    def start_bassin_thread(self):
        if not self.ze_shp_var.get().strip():
            messagebox.showerror("Erreur", "Veuillez d'abord sélectionner un shapefile pour la Zone d'étude.")
            return
        self.bassin_button.config(state="disabled")
        # Use shared browser if available
        driver = self._get_or_create_driver() if hasattr(self, '_get_or_create_driver') else None
        t = threading.Thread(target=self._run_bassin, args=(driver,))
        t.daemon = True
        t.start()



    def open_gmaps(self):

        if not self.ze_shp_var.get().strip():

            messagebox.showerror("Erreur", "Sélectionner la Zone d'étude.")

            return

        try:

            gdf = gpd.read_file(self.ze_shp_var.get())

            if gdf.crs is None:

                raise ValueError("CRS non défini")

            gdf = gdf.to_crs("EPSG:4326")

            centroid = gdf.geometry.unary_union.centroid

            lat, lon = centroid.y, centroid.x

            url = f"https://www.google.com/maps/@{lat},{lon},17z"

            print(f"[Maps] {url}", file=self.stdout_redirect)

            webbrowser.open(url)

        except Exception as e:

            messagebox.showerror("Erreur", f"Impossible d’ouvrir Google Maps : {e}")



    def start_rlt_thread(self):
        if not self.ze_shp_var.get().strip():
            messagebox.showerror("Erreur", "Sélectionner la Zone d'étude.")
            return
        self.rlt_button.config(state="disabled")
        # Use shared browser if available
        driver = self._get_or_create_driver() if hasattr(self, '_get_or_create_driver') else None
        t = threading.Thread(target=self._run_rlt, args=(driver,))
        t.daemon = True
        t.start()


    def _run_rlt(self, driver=None):
        try:
            ze_path = self.ze_shp_var.get()
            gdf = gpd.read_file(ze_path)
            if gdf.crs is None:
                raise ValueError("CRS non défini")
            gdf = gdf.to_crs("EPSG:4326")
            centroid = gdf.geometry.unary_union.centroid
            lat_dd, lon_dd = centroid.y, centroid.x
            commune, _dep = self._detect_commune(lat_dd, lon_dd)
            wait_s = WAIT_TILES_DEFAULT
            out_dir = OUTPUT_DIR_RLT
            os.makedirs(out_dir, exist_ok=True)
            comment_txt = COMMENT_TEMPLATE.format(commune=commune)

            # Use shared driver or create new one if not provided
            if driver is None:
                driver = self._get_or_create_driver()
                if not driver:
                    print("[IGN] Impossible de créer le navigateur", file=self.stdout_redirect)
                    return
                should_cleanup = True
            else:
                should_cleanup = False
                # Open in new tab
                driver.execute_script("window.open('');")
                driver.switch_to.window(driver.window_handles[-1])

            print(f"[IGN] Utilisation du navigateur partagé…", file=self.stdout_redirect)



            images = []

            viewport = (By.CSS_SELECTOR, "div.ol-viewport")

            for title, layer_val in LAYERS:

                url = URL.format(lon=f"{lon_dd:.6f}", lat=f"{lat_dd:.6f}", layer=layer_val)

                print(f"[IGN] {title} ? {url}", file=self.stdout_redirect)

                driver.get(url)

                WebDriverWait(driver, 20).until(EC.visibility_of_element_located(viewport))

                time.sleep(wait_s)

                tgt = driver.find_element(*viewport)

                img_path = os.path.join(out_dir, f"{title}.png")

                if tgt.screenshot(img_path):

                    img = Image.open(img_path)

                    w, h = img.size

                    left, right = int(w * 0.05), int(w * 0.95)

                    img.crop((left, 0, right, h)).save(img_path)

                    images.append((title, img_path))

                    print(f"[IGN] Capture OK : {img_path}", file=self.stdout_redirect)

                else:

                    print(f"[IGN] Capture échouée : {title}", file=self.stdout_redirect)



            driver.quit()



            if not images:

                print("[IGN] Aucune image ? pas de doc.", file=self.stdout_redirect)

                messagebox.showwarning("IGN", "Aucune image capturée.")

                return



            print("[IGN] Génération du Word…", file=self.stdout_redirect)

            doc = Document()

            style_normal = doc.styles['Normal']

            style_normal.font.name = 'Calibri'

            style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

            style_tbl = doc.styles['Table Grid']

            style_tbl.font.name = 'Calibri'

            style_tbl._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')



            sec = doc.sections[0]

            sec.orientation = WD_ORIENT.LANDSCAPE

            sec.page_width, sec.page_height = sec.page_height, sec.page_width

            for m in (sec.left_margin, sec.right_margin, sec.top_margin, sec.bottom_margin):

                m = Cm(1.5)



            cap_par = doc.add_paragraph()

            cap_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

            add_hyperlink(cap_par, "https://remonterletemps.ign.fr/",

                          f"Comparaison temporelle — {commune} (source : IGN – RemonterLeTemps)")



            table = doc.add_table(rows=2, cols=2)

            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            table.style = "Table Grid"

            table.autofit = False



            for idx, (title, path) in enumerate(images):

                r, c = divmod(idx, 2)

                cell = table.cell(r, c)

                p_t = cell.paragraphs[0]

                run_t = p_t.add_run(title); run_t.bold = True

                p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell.add_paragraph()

                p_img = cell.add_paragraph()

                if os.path.exists(path):

                    p_img.add_run().add_picture(path, width=IMG_WIDTH)

                else:

                    p_img.add_run(f"[image manquante : {title}]")

                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER



            doc.add_paragraph()

            p_comm = doc.add_paragraph(comment_txt)

            p_comm.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY



            doc_path = os.path.join(out_dir, WORD_FILENAME)

            doc.save(doc_path)

            print(f"[IGN] Document généré : {doc_path}", file=self.stdout_redirect)

            self._set_status(f"Terminé — {doc_path}")

        except Exception as e:

            print(f"[IGN] Erreur : {e}", file=self.stdout_redirect)

            messagebox.showerror("IGN", str(e))

        finally:

            self.after(0, lambda: self.rlt_button.config(state="normal"))



    def _run_bassin(self, driver=None):
        try:

            ze_path = self.ze_shp_var.get()

            gdf = gpd.read_file(ze_path)

            if gdf.crs is None:

                raise ValueError("CRS non défini")

            gdf = gdf.to_crs("EPSG:4326")

            centroid = gdf.geometry.unary_union.centroid

            lat_dd, lon_dd = centroid.y, centroid.x

            user_address = dd_to_dms(lat_dd, lon_dd)

            download_dir = OUT_IMG

            target_folder_name = "Bassin versant"

            target_path = os.path.join(download_dir, target_folder_name)

            os.makedirs(download_dir, exist_ok=True)

            print(f"[BV] Coordonnées : {lat_dd:.6f}, {lon_dd:.6f}", file=self.stdout_redirect)

            # Use shared driver or create new one if not provided
            if driver is None:
                # Create driver with download preferences for this specific task
                options = webdriver.ChromeOptions()
                options.add_argument("--log-level=3")
                options.add_experimental_option('excludeSwitches', ['enable-logging'])
                options.add_argument("--disable-extensions")
                prefs = {
                    "download.default_directory": download_dir,
                    "download.prompt_for_download": False,
                    "profile.default_content_settings.popups": 0,
                    "download.directory_upgrade": True
                }
                options.add_experimental_option("prefs", prefs)
                driver = webdriver.Chrome(options=options)
                try:
                    driver.maximize_window()
                except Exception:
                    pass
                should_cleanup = True
                print("[BV] Navigateur créé avec préférences de téléchargement...", file=self.stdout_redirect)
            else:
                should_cleanup = False
                # Open in new tab
                driver.execute_script("window.open('');")
                driver.switch_to.window(driver.window_handles[-1])
                print("[BV] Utilisation du navigateur partagé (nouvel onglet)...", file=self.stdout_redirect)

            url = "https://mghydro.com/watersheds/"

            print(f"[BV] Navigation vers {url}...", file=self.stdout_redirect)

            driver.get(url)

            wait = WebDriverWait(driver, 2)

            opts_button = wait.until(EC.element_to_be_clickable((By.ID, "opts_click")))

            opts_button.click(); time.sleep(1)

            downloadable_checkbox = wait.until(EC.element_to_be_clickable((By.ID, "downloadable")))

            if not downloadable_checkbox.is_selected():

                downloadable_checkbox.click()

            search_icon = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, ".leaflet-control-search .search-button")))

            search_icon.click()

            search_input = wait.until(EC.visibility_of_element_located((By.ID, "searchtext84")))

            search_input.clear(); time.sleep(0.3)

            search_input.send_keys(user_address); time.sleep(1.5)

            search_input.send_keys(Keys.ARROW_DOWN); time.sleep(0.3)

            search_input.send_keys(Keys.ENTER); time.sleep(1.5)

            map_element = wait.until(EC.presence_of_element_located((By.ID, "map")))

            ActionChains(driver).move_to_element(map_element).click().perform(); time.sleep(0.8)

            wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "div.leaflet-popup")))

            delineate_button = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, ".leaflet-popup .gobutton")))

            delineate_button.click(); time.sleep(1.5)

            downloads_button = wait.until(EC.element_to_be_clickable((By.XPATH, "//span[contains(@class, 'ui-selectmenu-text') and contains(text(), 'Watershed Boundary')]")))

            downloads_button.click(); time.sleep(0.8)

            ActionChains(driver).send_keys(Keys.ARROW_DOWN).pause(0.3).send_keys(Keys.ARROW_DOWN).pause(0.3).send_keys(Keys.ENTER).perform()

            time.sleep(1.5)

        except Exception as e:

            print(f"[BV] Erreur Selenium : {e}", file=self.stdout_redirect)

        finally:

            # Only cleanup if we created the driver ourselves
            if should_cleanup:
                try:
                    driver.quit()
                except Exception:
                    pass
            else:
                # Just close the current tab
                try:
                    driver.close()
                    if len(driver.window_handles) > 0:
                        driver.switch_to.window(driver.window_handles[0])
                except Exception:
                    pass



        try:

            print("[BV] Attente du téléchargement du ZIP...", file=self.stdout_redirect)

            zip_file_path = None

            wait_time = 30

            start_time = time.time()

            while time.time() - start_time < wait_time:

                zip_candidates = [f for f in os.listdir(download_dir) if f.lower().endswith('.zip') and not f.lower().endswith('.crdownload')]

                if zip_candidates:

                    zip_candidates_full = [os.path.join(download_dir, z) for z in zip_candidates]

                    zip_file_path = max(zip_candidates_full, key=os.path.getmtime)

                    size1 = os.path.getsize(zip_file_path); time.sleep(1); size2 = os.path.getsize(zip_file_path)

                    if size1 == size2:

                        print(f"[BV] ZIP détecté : {os.path.basename(zip_file_path)}", file=self.stdout_redirect)

                        break

                time.sleep(1)



            if not zip_file_path:

                print("[BV] Aucun fichier ZIP trouvé.", file=self.stdout_redirect)

            else:

                if os.path.exists(target_path):

                    print(f"[BV] Remplacement du dossier '{target_folder_name}'", file=self.stdout_redirect)

                    shutil.rmtree(target_path, ignore_errors=True)

                os.makedirs(target_path, exist_ok=True)

                with zipfile.ZipFile(zip_file_path, 'r') as zf:

                    zf.extractall(path=target_path)

                os.remove(zip_file_path)

                print(f"[BV] Décompression terminée dans '{target_folder_name}'", file=self.stdout_redirect)

                self._set_status(f"Bassin versant : {target_path}")

        except Exception as e:

            print(f"[BV] Erreur décompression : {e}", file=self.stdout_redirect)

        finally:

            self.after(0, lambda: self.bassin_button.config(state="normal"))



    def _set_status(self, txt: str):

        self.after(0, lambda: self.status_label.config(text=txt))



    def _detect_commune(self, lat: float, lon: float) -> Tuple[str, str]:

        try:

            url = ("https://nominatim.openstreetmap.org/reverse?format=json"

                   f"&lat={lat}&lon={lon}&zoom=10&addressdetails=1")

            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})

            with urllib.request.urlopen(req, timeout=10) as resp:

                data = json.load(resp)

            addr = data.get("address", {})

            commune = "Inconnue"

            for key in ("city", "town", "village", "municipality"):

                if key in addr:

                    commune = addr[key]

                    break

            dept_name = addr.get("county", "")

            dep_rev = {v: k for k, v in DEP.items()}

            dep = dep_rev.get(dept_name, "")

            if not dep:

                postcode = addr.get("postcode", "")

                if len(postcode) >= 2:

                    dep = postcode[:2]

            return commune, dep

        except Exception as e:

            print(f"[Wiki] Détection commune échouée : {e}", file=self.stdout_redirect)

            return "Inconnue", ""



    # ---------- Gestion projets QGIS ----------

    def _populate_projects(self):

        for w in list(self.scrollable_frame.children.values()): w.destroy()

        self.project_vars = {}

        self.all_projects = discover_projects()

        self.filtered_projects = list(self.all_projects)

        if not self.all_projects:

            ttk.Label(self.scrollable_frame, text="Aucun projet trouvé ou dossier inaccessible.", foreground="red").pack(anchor="w")

            return

        # Display checkboxes in two columns for compactness

        for i, proj_path in enumerate(self.filtered_projects):

            var = tk.IntVar(value=1); self.project_vars[proj_path] = var

            r, c = divmod(i, 2)

            cb = ttk.Checkbutton(self.scrollable_frame, text=os.path.basename(proj_path), variable=var, style="Card.TCheckbutton")

            cb.grid(row=r, column=c, sticky='w', padx=4, pady=2)

        try:

            self.scrollable_frame.columnconfigure(0, weight=1)

            self.scrollable_frame.columnconfigure(1, weight=1)

        except Exception:

            pass



    def _apply_filter(self):

        term = normalize_name(self.filter_var.get())

        for w in list(self.scrollable_frame.children.values()): w.destroy()

        self.filtered_projects = [p for p in self.all_projects if term in normalize_name(os.path.basename(p))]

        if not self.filtered_projects:

            ttk.Label(self.scrollable_frame, text="Aucun projet ne correspond au filtre.", foreground="red").pack(anchor="w")

            self.project_vars = {}; self._update_counts(); return

        for i, proj_path in enumerate(self.filtered_projects):

            current = self.project_vars.get(proj_path, tk.IntVar(value=1))

            self.project_vars[proj_path] = current

            r, c = divmod(i, 2)

            cb = ttk.Checkbutton(self.scrollable_frame, text=os.path.basename(proj_path), variable=current, style="Card.TCheckbutton")

            cb.grid(row=r, column=c, sticky='w', padx=4, pady=2)

        try:

            self.scrollable_frame.columnconfigure(0, weight=1)

            self.scrollable_frame.columnconfigure(1, weight=1)

        except Exception:

            pass

        self._update_counts()



    def _select_all(self, state: bool):

        for var in self.project_vars.values(): var.set(1 if state else 0)

        self._update_counts()



    def _selected_projects(self) -> List[str]:

        return [p for p, v in self.project_vars.items() if v.get() == 1 and p in self.filtered_projects]



    def _update_counts(self):

        selected = len(self._selected_projects()); total = len(self.filtered_projects)

        self.status_label.config(text=f"Projets sélectionnés : {selected} / {total}")



    # ---------- Lancement export ----------

    def start_export_thread(self):

        if self.busy:

            print("Une action est déjà en cours.", file=self.stdout_redirect)

            return

        if not self.ze_shp_var.get() or not self.ae_shp_var.get():

            messagebox.showerror("Erreur", "Sélectionnez les deux shapefiles."); return

        if not os.path.isfile(self.ze_shp_var.get()) or not os.path.isfile(self.ae_shp_var.get()):

            messagebox.showerror("Erreur", "Un shapefile est introuvable."); return

        projets = self._selected_projects()

        if not projets:

            messagebox.showerror("Erreur", "Sélectionnez au moins un projet."); return



        self.busy = True

        self.export_button.config(state="disabled")
        try:
            if hasattr(self, 'id_button') and self.id_button:
                self.id_button.config(state="disabled")
        except Exception:
            pass

        mode = self.cadrage_var.get()

        exp_type = self.export_type_var.get()

        png_exports = 2 if (exp_type in ("PNG", "BOTH") and mode == "BOTH") else (1 if exp_type in ("PNG", "BOTH") else 0)

        qgs_exports = 1 if exp_type in ("QGS", "BOTH") else 0

        per_project = png_exports + qgs_exports

        self.total_expected = per_project * len(projets)

        self.progress_done = 0

        self.progress.config(mode="determinate", maximum=max(1, self.total_expected), value=0)

        self.status_label.config(text=f"Progression : 0/{self.total_expected}")



        self.prefs.update({

            "ZE_SHP": self.ze_shp_var.get(),

            "AE_SHP": self.ae_shp_var.get(),

            "CADRAGE_MODE": self.cadrage_var.get(),

            "OVERWRITE": bool(self.overwrite_var.get()),

            "DPI": int(self.dpi_var.get()),

            "N_WORKERS": int(self.workers_var.get()),

            "MARGIN_FAC": float(self.margin_var.get()),

            "OUT_DIR": self.out_dir_var.get(),

            "EXPORT_TYPE": exp_type,

        }); save_prefs(self.prefs)



        t = threading.Thread(target=self._run_export_logic, args=(projets,), daemon=True)

        t.start()



    def _run_export_logic(self, projets: List[str]):

        old_stdout = sys.stdout

        sys.stdout = self.stdout_redirect

        try:

            start = datetime.datetime.now()

            out_dir = self.out_dir_var.get() or OUT_IMG

            os.makedirs(out_dir, exist_ok=True)

            log_with_time(f"{len(projets)} projets (attendu = calcul en cours)")

            log_with_time(f"Workers={self.workers_var.get()}, DPI={self.dpi_var.get()}, marge={self.margin_var.get():.2f}, overwrite={self.overwrite_var.get()}")

            workers = int(self.workers_var.get())

            # Désactive provisoirement le multiprocessing pour éviter les erreurs _multiprocessing

            workers = workers

            chunks = chunk_even(projets, workers)

            # Forcer au moins 2 workers pour utiliser ProcessPoolExecutor

            # (et donc le Python de QGIS configuré ci-dessous)

            workers = max(1, workers)

            cfg = {

                "QGIS_ROOT": QGIS_ROOT,

                "QGIS_APP": QGIS_APP,

                "PY_VER": PY_VER,

                "EXPORT_DIR": out_dir,

                "DPI": int(self.dpi_var.get()),

                "MARGIN_FAC": float(self.margin_var.get()),

                "LAYER_AE_NAME": LAYER_AE_NAME,

                "LAYER_ZE_NAME": LAYER_ZE_NAME,

                "AE_SHP": self.ae_shp_var.get(),

                "ZE_SHP": self.ze_shp_var.get(),

                "CADRAGE_MODE": self.cadrage_var.get(),

                "OVERWRITE": bool(self.overwrite_var.get()),

                "EXPORT_TYPE": self.export_type_var.get(),

                "WORKERS": workers,

            }

            ok_total = 0

            ko_total = 0

            def ui_update_progress(done_inc):

                self.progress_done += done_inc

                self.progress["value"] = min(self.progress_done, self.total_expected)

                self.status_label.config(text=f"Progression : {self.progress_done}/{self.total_expected}")

            if workers <= 1:

                for chunk in chunks:

                    ok, ko = run_worker_subprocess(chunk, cfg)

                    ok_total += ok

                    ko_total += ko

                    self.after(0, ui_update_progress, ok + ko)

                    log_with_time(f"Lot terminé: {ok} OK, {ko} KO")

            else:

                try:

                    import multiprocessing as mp

                    # Nettoyage de l'environnement hérité pour éviter collisions Python 3.12/3.13

                    for _k in ("PYTHONHOME", "PYTHONPATH", "PYTHONSTARTUP"):

                        try:

                            os.environ.pop(_k, None)

                        except Exception:

                            pass

                    os.environ["PYTHONNOUSERSITE"] = "1"

                    # Fixer PYTHONHOME et PYTHONPATH sur le Python de QGIS

                    try:

                        qgis_py_root = os.path.join(QGIS_ROOT, "apps", PY_VER)

                        if os.path.isdir(qgis_py_root):

                            os.environ["PYTHONHOME"] = qgis_py_root

                            qgis_lib   = os.path.join(qgis_py_root, "Lib")

                            qgis_dlls  = os.path.join(qgis_py_root, "DLLs")

                            qgis_site  = os.path.join(qgis_lib, "site-packages")

                            qgis_app_py = os.path.join(QGIS_APP, "python")

                            py_paths = [qgis_py_root, qgis_lib, qgis_dlls, qgis_site, qgis_app_py]

                            os.environ["PYTHONPATH"] = os.pathsep.join(py_paths)

                            # Préfixer le PATH avec les dossiers Python QGIS pour la résolution des DLLs

                            os.environ["PATH"] = os.pathsep.join([qgis_py_root, qgis_dlls, os.environ.get("PATH", "")])

                            log_with_time(f"PYTHONHOME={qgis_py_root}")

                    except Exception:

                        pass

                    # Fixer PYTHONHOME sur le Python de QGIS pour que l'interprète trouve sa stdlib

                    try:

                        qgis_py_root = os.path.join(QGIS_ROOT, "apps", PY_VER)

                        if os.path.isdir(qgis_py_root):

                            os.environ["PYTHONHOME"] = qgis_py_root

                            log_with_time(f"PYTHONHOME={qgis_py_root}")

                    except Exception:

                        pass

                    ctx = mp.get_context("spawn")

                    try:

                        qgis_py = os.path.join(QGIS_ROOT, "apps", PY_VER, "python.exe")

                        if os.path.isfile(qgis_py):

                            ctx.set_executable(qgis_py)

                            log_with_time(f"MP exe: {qgis_py}")

                        else:

                            log_with_time(f"Python QGIS introuvable: {qgis_py}")

                    except Exception as e:

                        log_with_time(f"set_executable échec: {e}")

                except Exception as e:

                    log_with_time(f"init multiprocessing: {e}")

                # Ajuster temporairement sys.path pour privilégier les libs QGIS

                old_syspath = list(sys.path)

                try:

                    qgis_py_root = os.path.join(QGIS_ROOT, "apps", PY_VER)

                    qgis_py_lib = os.path.join(qgis_py_root, "Lib")

                    qgis_site = os.path.join(qgis_py_lib, "site-packages")

                    qgis_app_py = os.path.join(QGIS_APP, "python")

                    def _keep_path(p: str) -> bool:

                        if not isinstance(p, str):

                            return False

                        l = p.lower()

                        if "python313" in l or "python311" in l or "python310" in l or "python39" in l:

                            return False

                        if ".venv" in l:

                            return False

                        return True

                    sys.path = [qgis_py_root, qgis_py_lib, qgis_site, qgis_app_py] + [p for p in old_syspath if _keep_path(p)]

                except Exception as e:

                    log_with_time(f"sys.path cleanup skip: {e}")

                with ThreadPoolExecutor(max_workers=workers) as ex:

                    futures = [ex.submit(run_worker_subprocess, chunk, cfg) for chunk in chunks if chunk]

                    for fut in as_completed(futures):

                        try:

                            ok, ko = fut.result()

                            ok_total += ok

                            ko_total += ko

                            self.after(0, ui_update_progress, ok + ko)

                            log_with_time(f"Lot terminé: {ok} OK, {ko} KO")

                        except Exception as e:

                            log_with_time(f"Erreur worker: {e}")

                # Restaure le sys.path initial

                sys.path = old_syspath

            elapsed = datetime.datetime.now() - start

            log_with_time(f"FIN — OK={ok_total} | KO={ko_total} | Attendu={self.total_expected} | Durée={elapsed}")

            self.after(0, lambda: self.status_label.config(text=f"Terminé — OK={ok_total} / KO={ko_total}"))

        except Exception as e:

            log_with_time(f"Erreur critique: {e}")

            _err = str(e)

            self.after(0, lambda msg=_err: messagebox.showerror("Erreur", msg))

        finally:

            sys.stdout = old_stdout

            self.after(0, self._run_finished)



    # ---------- Lancement ID contexte ----------

    def start_id_thread(self):

        if self.busy:

            print("Une action est déjà en cours.", file=self.stdout_redirect)

            return

        ae = to_long_unc(os.path.normpath(self.ae_shp_var.get().strip()))

        ze = to_long_unc(os.path.normpath(self.ze_shp_var.get().strip()))

        if not ae or not ze:

            messagebox.showerror("Erreur", "Sélectionnez les deux shapefiles."); return

        if not os.path.isfile(ae) or not os.path.isfile(ze):

            messagebox.showerror("Erreur", "Un shapefile est introuvable."); return



        self.busy = True

        self.export_button.config(state="disabled")
        # Certaines configurations n'affichent pas le bouton d'identification
        try:
            if hasattr(self, 'id_button') and self.id_button:
                self.id_button.config(state="disabled")
        except Exception:
            pass

        self.progress.config(mode="indeterminate")

        self.progress.start()

        self.status_label.config(text="Analyse en cours…")



        # Enregistrer des chemins standard (sans préfixe long UNC) dans les préférences
        self.prefs.update({
            "ZE_SHP": from_long_unc(ze),
            "AE_SHP": from_long_unc(ae),
            "ID_TAMPON_KM": float(self.buffer_var.get()),
        }); save_prefs(self.prefs)



        t = threading.Thread(target=self._run_id_logic, args=(ae, ze, float(self.buffer_var.get())), daemon=True)

        t.start()



    def _run_id_logic(self, ae: str, ze: str, buffer_km: float):

        old_stdout = sys.stdout

        sys.stdout = self.stdout_redirect

        try:

            from .id_contexte_eco import run_analysis as run_id_context

            run_id_context(ae, ze, buffer_km)

            log_with_time("Analyse terminée.")

            self.after(0, lambda: self.status_label.config(text="Terminé"))

        except Exception as e:

            log_with_time(f"Erreur: {e}")

            _err = str(e)

            self.after(0, lambda msg=_err: messagebox.showerror("Erreur", msg))

        finally:

            sys.stdout = old_stdout

            self.after(0, lambda: (self.progress.stop(), self.progress.config(mode="determinate", value=0)))

            self.after(0, self._run_finished)



    def _run_finished(self):

        self.export_button.config(state="normal")
        try:
            if hasattr(self, 'id_button') and self.id_button:
                self.id_button.config(state="normal")
        except Exception:
            pass

        self.busy = False



# =========================
# Onglets PlantNet et Biodiv (stubs)
# =========================
class PlantNetTab(ttk.Frame):
    def __init__(self, parent, style_helper, prefs):
        super().__init__(parent)
        self.style_helper = style_helper
        self.prefs = prefs
        
        # Simple placeholder UI
        ttk.Label(self, text="Pl@ntNet - Identification des plantes", 
                 font=("Segoe UI", 15, "bold")).pack(pady=20)
        ttk.Label(self, text="Fonctionnalité à venir...", 
                 style="Card.TLabel").pack(pady=10)

class BiodivTab(ttk.Frame):
    def __init__(self, parent, style_helper, prefs):
        super().__init__(parent)
        self.style_helper = style_helper
        self.prefs = prefs
        
        # Simple placeholder UI
        ttk.Label(self, text="Biodiv'AURA - Base de données", 
                 font=("Segoe UI", 15, "bold")).pack(pady=20)
        ttk.Label(self, text="Fonctionnalité à venir...", 
                 style="Card.TLabel").pack(pady=10)

# =========================
# App principale avec Notebook
# =========================
class Application(tk.Tk):
    def __init__(self):
        super().__init__()
        self.prefs = load_prefs()

        self.title("Bota-Logiciel | Assistant Contexte Éco & Identification")
        self.geometry(self.prefs.get("GEOMETRY", "1200x900"))

        self.style_helper = StyleHelper(self, self.prefs)
        self.style_helper.apply(self.prefs.get("THEME", "light"))

        self.notebook = ttk.Notebook(self)
        self.export_tab = ExportCartesTab(self.notebook, self.style_helper, self.prefs)
        self.plantnet_tab = PlantNetTab(self.notebook, self.style_helper, self.prefs)
        self.biodiv_tab = BiodivTab(self.notebook, self.style_helper, self.prefs)

        self.notebook.add(self.export_tab, text="Contexte Écologique & Cartes")
        self.notebook.add(self.plantnet_tab, text="  Pl@ntNet  ")
        self.notebook.add(self.biodiv_tab, text="Biodiv'AURA")

        self.notebook.pack(expand=True, fill='both', padx=10, pady=10)

        self.protocol("WM_DELETE_WINDOW", self._on_closing)

    def _on_closing(self):
        try:
            g = self.geometry()
            if "x" in g:
                self.prefs["GEOMETRY"] = g
            save_prefs(self.prefs)
        except Exception:
            pass

        # Call the cleanup method for the shared browser
        if hasattr(self, 'export_tab'):
            self.export_tab._cleanup_driver()

        self.destroy()


def launch():
    """Point d'entrée programmatique pour lancer l'application GUI.
    Utilisé par `Start.py` via `modules.main_app.launch()`.
    """
    app = Application()
    app.mainloop()


if __name__ == '__main__':
    launch()
















