#!/usr/bin/env "C:/Program Files/QGIS 3.40.3/apps/Python312/python.exe"
# -*- coding: utf-8 -*-
r"""
Application à onglets (ttk.Notebook) :

Onglet 1 — « Export Cartes » :
    - Identique fonctionnellement à la version précédente (UI modernisée).
    - Exporte les PNG des mises en page QGIS avec cadrage AE/ZE.
    - Parallélisation, logs, thème clair/sombre, préférences.

Onglet 2 — « Remonter le temps » :
    - Reprend le workflow IGN (capture + rapport Word) mais
      DEMANDE à l’utilisateur :
        • Coordonnées en DMS (LAT puis LON, ex : 45°09'30" N 5°43'12" E)
        • Nom de la commune
      au lieu de lire un Excel.
    - Options : dossier de sortie, headless, tempo de chargement.
    - Produit un Word 2×2 avec les vues temporelles, + commentaire.

Pré-requis Python : qgis (environnement QGIS), selenium, pillow, python-docx, openpyxl (non utilisé ici), chromedriver dans PATH.
"""

import os
import sys
import re
import json
import time
import shutil
import tempfile
import datetime
import threading
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter import font as tkfont
from typing import List, Optional, Tuple
from concurrent.futures import ProcessPoolExecutor, as_completed

# ==== Imports spécifiques onglet 2 (gardés en tête de fichier comme le script source) ====
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from PIL import Image
import requests
from io import BytesIO
import pillow_heif  # Importer pillow-heif pour la prise en charge des fichiers HEIC

# Enregistrer le décodeur HEIF
pillow_heif.register_heif_opener()

# =========================
# Paramètres globaux
# =========================
# Onglet 1 — Export Cartes (inchangé)
DPI_DEFAULT        = 300
N_WORKERS_DEFAULT  = max(1, min((os.cpu_count() or 2) - 1, 6))
MARGIN_FAC_DEFAULT = 1.15
OVERWRITE_DEFAULT  = False

LAYER_AE_NAME = "Aire d'étude élargie"
LAYER_ZE_NAME = "Zone d'étude"

BASE_SHARE = r"\\192.168.1.240\commun\PARTAGE"
SUBPATH    = r"Espace_RWO\CARTO ROBIN"

OUT_IMG    = r"C:\Users\utilisateur\Mon Drive\1 - Bota & Travail\+++++++++  BOTA  +++++++++\---------------------- 3) BDD\PYTHON\2) Contexte éco\OUTPUT"

# Dossier par défaut pour la sélection des shapefiles (onglet 1)
DEFAULT_SHAPE_DIR = r"C:\Users\utilisateur\Mon Drive\1 - Bota & Travail\+++++++++  BOTA  +++++++++\---------------------- 2) CARTO terrain"

# QGIS
QGIS_ROOT = r"C:\Program Files\QGIS 3.40.3"
QGIS_APP  = os.path.join(QGIS_ROOT, "apps", "qgis")
PY_VER    = "Python312"

# Préférences
PREFS_PATH = os.path.join(os.path.expanduser("~"), "ExportCartesContexteEco.config.json")

# Onglet 2 — Remonter le temps (constants issus du script source)
LAYERS = [
    ("Aujourd’hui",   "10"),
    ("2000-2005",     "18"),
    ("1965-1980",     "20"),
    ("1950-1965",     "19"),
]
URL = ("https://remonterletemps.ign.fr/comparer/?lon={lon}&lat={lat}"
       "&z=17&layer1={layer}&layer2=19&mode=dub1")
WAIT_TILES_DEFAULT = 1.5
IMG_WIDTH = Cm(12.5 * 0.8)
WORD_FILENAME = "Comparaison_temporelle_Paysage.docx"
OUTPUT_DIR_RLT = os.path.join(OUT_IMG, "Remonter le temps")
COMMENT_TEMPLATE = (
    "Rédige un commentaire synthétique de l'évolution de l'occupation du sol observée "
    "sur les images aériennes de la zone d'étude, aux différentes dates indiquées "
    "(1950–1965, 1965–1980, 2000–2005, aujourd’hui). Concentre-toi sur les grandes "
    "dynamiques d'aménagement (urbanisation, artificialisation, évolution des milieux "
    "ouverts ou boisés), en identifiant les principales transformations visibles. "
    "Fais ta réponse en un seul court paragraphe. Intègre les éléments de contexte "
    "historique et territorial propres à la commune de {commune} pour interpréter ces évolutions."
)

# =========================
# Utils communs
# =========================
def log_with_time(msg: str) -> None:
    print(f"[{datetime.datetime.now().strftime('%H:%M:%S')}] {msg}")

def normalize_name(s: str) -> str:
    s2 = s.replace("\u00A0", " ").replace("\u202F", " ")
    while "  " in s2:
        s2 = s2.replace("  ", " ")
    return s2.strip().lower()

def to_long_unc(path: str) -> str:
    if path.startswith("\\\\?\\"): return path
    if path.startswith("\\\\"):   return "\\\\?\\UNC" + path[1:]
    return "\\\\?\\" + path

def chunk_even(lst: List[str], k: int) -> List[List[str]]:
    if not lst: return []
    k = max(1, min(k, len(lst)))
    base = len(lst) // k
    extra = len(lst) % k
    out, start = [], 0
    for i in range(k):
        size = base + (1 if i < extra else 0)
        out.append(lst[start:start+size])
        start += size
    return out

def load_prefs() -> dict:
    if os.path.isfile(PREFS_PATH):
        try:
            with open(PREFS_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}

def save_prefs(d: dict) -> None:
    try:
        with open(PREFS_PATH, "w", encoding="utf-8") as f:
            json.dump(d, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

class TextRedirector:
    def __init__(self, widget): self.widget = widget
    def write(self, s):
        self.widget.config(state='normal')
        self.widget.insert(tk.END, s)
        self.widget.see(tk.END)
        self.widget.config(state='disabled')
        self.widget.update_idletasks()
    def flush(self): pass

class ToolTip:
    def __init__(self, widget, text: str, delay: int = 600):
        self.widget = widget; self.text = text; self.delay = delay
        self.tipwindow = None; self.id = None
        widget.bind("<Enter>", self._schedule); widget.bind("<Leave>", self._hide)
    def _schedule(self, _=None):
        self._cancel(); self.id = self.widget.after(self.delay, self._show)
    def _cancel(self):
        if self.id: self.widget.after_cancel(self.id); self.id = None
    def _show(self):
        if self.tipwindow: return
        x = self.widget.winfo_rootx() + 15; y = self.widget.winfo_rooty() + 25
        self.tipwindow = tw = tk.Toplevel(self.widget); tw.wm_overrideredirect(True); tw.wm_geometry(f"+{x}+{y}")
        ttk.Label(tw, text=self.text, style="Tooltip.TLabel", padding=(8, 4)).pack()
    def _hide(self, _=None):
        self._cancel()
        if self.tipwindow: self.tipwindow.destroy(); self.tipwindow = None

# =========================
# Fonctions QGIS (onglet 1) — inchangées
# =========================
def worker_run(args: Tuple[List[str], dict]) -> Tuple[int, int]:
    projects, cfg = args

    os.environ["OSGEO4W_ROOT"] = cfg["QGIS_ROOT"]
    os.environ["QGIS_PREFIX_PATH"] = cfg["QGIS_APP"]
    os.environ.setdefault("GDAL_DATA", os.path.join(cfg["QGIS_ROOT"], "share", "gdal"))
    os.environ.setdefault("PROJ_LIB", os.path.join(cfg["QGIS_ROOT"], "share", "proj"))
    os.environ.setdefault("QT_QPA_FONTDIR", r"C:\Windows\Fonts")

    qt_base = None
    for name in ("Qt6", "Qt5"):
        base = os.path.join(cfg["QGIS_ROOT"], "apps", name)
        if os.path.isdir(base):
            qt_base = base
            break
    if qt_base is None:
        raise RuntimeError("Qt introuvable")

    platform_dir = os.path.join(qt_base, "plugins", "platforms")
    os.environ["QT_PLUGIN_PATH"] = os.path.join(qt_base, "plugins")
    os.environ["QT_QPA_PLATFORM_PLUGIN_PATH"] = platform_dir
    qpa = "windows" if os.path.isfile(os.path.join(platform_dir, "qwindows.dll")) \
        else ("minimal" if os.path.isfile(os.path.join(platform_dir, "qminimal.dll")) else "offscreen")
    os.environ["QT_QPA_PLATFORM"] = qpa

    os.environ["PATH"] = os.pathsep.join([
        os.path.join(qt_base, "bin"),
        os.path.join(cfg["QGIS_APP"], "bin"),
        os.path.join(cfg["QGIS_ROOT"], "bin"),
        os.environ.get("PATH", ""),
    ])

    sys.path.insert(0, os.path.join(cfg["QGIS_APP"], "python"))
    sys.path.insert(0, os.path.join(cfg["QGIS_ROOT"], "apps", cfg["PY_VER"], "Lib", "site-packages"))

    from qgis.core import (
        QgsApplication, QgsProject, QgsLayoutExporter, QgsLayoutItemMap, QgsRectangle,
        QgsCoordinateTransform
    )

    qgs = QgsApplication([], False)
    qgs.setPrefixPath(cfg["QGIS_APP"], True)
    qgs.initQgis()

    ok = 0
    ko = 0

    def adjust_extent_to_item_ratio(ext: 'QgsRectangle', target_ratio: float, margin: float) -> 'QgsRectangle':
        if ext.width() <= 0 or ext.height() <= 0:
            return ext
        cx, cy = ext.center().x(), ext.center().y()
        w, h = ext.width(), ext.height()
        if (w / h) > target_ratio:
            new_h = w / target_ratio
            dh = (new_h - h) / 2.0
            xmin, xmax = ext.xMinimum(), ext.xMaximum()
            ymin, ymax = cy - h/2.0 - dh, cy + h/2.0 + dh
        else:
            new_w = h * target_ratio
            dw = (new_w - w) / 2.0
            ymin, ymax = ext.yMinimum(), ext.yMaximum()
            xmin, xmax = cx - w/2.0 - dw, cx + w/2.0 + dw
        cw, ch = (xmax - xmin), (ymax - ymin)
        mx, my = (margin - 1.0) * cw / 2.0, (margin - 1.0) * ch / 2.0
        return QgsRectangle(xmin - mx, ymin - my, xmax + mx, ymax + my)

    def extent_in_project_crs(prj: 'QgsProject', lyr) -> Optional['QgsRectangle']:
        ext = lyr.extent()
        try:
            if lyr.crs() != prj.crs():
                ct = QgsCoordinateTransform(lyr.crs(), prj.crs(), prj)
                ext = ct.transformBoundingBox(ext)
        except Exception:
            pass
        return ext

    def apply_extent_and_export(layout, lyr_extent: 'QgsRectangle', out_png: str) -> bool:
        maps = [it for it in layout.items() if isinstance(it, QgsLayoutItemMap)]
        if not maps: return False
        for m in maps:
            size = m.sizeWithUnits()
            target_ratio = max(1e-9, float(size.width()) / float(size.height()))
            adj_extent = adjust_extent_to_item_ratio(lyr_extent, target_ratio, cfg["MARGIN_FAC"])
            m.setExtent(adj_extent); m.refresh()
        img = QgsLayoutExporter.ImageExportSettings()
        img.dpi = cfg["DPI"]
        for attr in ("antialiasing", "antiAliasing"):
            if hasattr(img, attr):
                setattr(img, attr, True)
        try:
            flag_val = 0
            for name in dir(img.__class__):
                if "UseAdvancedEffects" in name:
                    flag_val |= int(getattr(img.__class__, name))
            if flag_val and hasattr(img, "flags"): img.flags = flag_val
        except Exception:
            pass
        if hasattr(img, "generateWorldFile"): img.generateWorldFile = False
        exp = QgsLayoutExporter(layout)
        res = exp.exportToImage(out_png, img)
        return res == QgsLayoutExporter.Success

    def relink_layer(prj: 'QgsProject', layer_name: str, shp_path: str) -> Optional[object]:
        layers = prj.mapLayersByName(layer_name)
        if not layers: return None
        lyr = layers[0]
        try:
            lyr.setDataSource(shp_path, lyr.name(), "ogr")
            return lyr
        except Exception:
            return None

    def export_views(projet_path: str) -> Tuple[int, int]:
        okc = 0; koc = 0
        nom = os.path.splitext(os.path.basename(projet_path))[0]
        out_ae = os.path.join(cfg["OUT_IMG"], f"{nom}__AE.png")
        out_ze = os.path.join(cfg["OUT_IMG"], f"{nom}__ZE.png")

        mode = cfg.get("CADRAGE_MODE", "BOTH")
        expected_exports = (1 if mode in ("AE", "BOTH") else 0) + (1 if mode in ("ZE", "BOTH") else 0)

        prj = QgsProject.instance(); prj.clear()

        opened = False
        for pth in (projet_path, to_long_unc(projet_path)):
            try:
                if prj.read(pth): opened = True; break
            except Exception:
                pass
        if not opened:
            return 0, expected_exports

        lm = prj.layoutManager()
        layouts = lm.layouts()
        if not layouts:
            prj.clear(); return 0, expected_exports
        layout = layouts[0]

        lyr_ae = relink_layer(prj, cfg["LAYER_AE_NAME"], cfg["AE_SHP"])
        lyr_ze = relink_layer(prj, cfg["LAYER_ZE_NAME"], cfg["ZE_SHP"])

        if mode in ("AE", "BOTH"):
            if (not cfg["OVERWRITE"]) and os.path.exists(out_ae):
                okc += 1
            else:
                if lyr_ae:
                    ext_ae = extent_in_project_crs(prj, lyr_ae)
                    if ext_ae and apply_extent_and_export(layout, ext_ae, out_ae): okc += 1
                    else: koc += 1
                else:
                    koc += 1

        if mode in ("ZE", "BOTH"):
            if (not cfg["OVERWRITE"]) and os.path.exists(out_ze):
                okc += 1
            else:
                if lyr_ze:
                    ext_ze = extent_in_project_crs(prj, lyr_ze)
                    if ext_ze and apply_extent_and_export(layout, ext_ze, out_ze): okc += 1
                    else: koc += 1
                else:
                    koc += 1

        prj.clear()
        return okc, koc

    for p in projects:
        ok_c, ko_c = export_views(p)
        ok += ok_c; ko += ko_c

    qgs.exitQgis()
    return ok, ko

def discover_projects() -> List[str]:
    partage = BASE_SHARE
    base_dir: Optional[str] = None
    try:
        if os.path.isdir(partage):
            for e in os.listdir(partage):
                if normalize_name(e) == normalize_name("Espace_RWO"):
                    base_espace = os.path.join(partage, e)
                    for s in os.listdir(base_espace):
                        if normalize_name(s) == normalize_name("CARTO ROBIN"):
                            base_dir = os.path.join(base_espace, s); break
                    break
    except Exception as e:
        log_with_time(f"Accès PARTAGE impossible via listdir: {e}")

    if not base_dir or not os.path.isdir(base_dir):
        base_dir = os.path.join(BASE_SHARE, SUBPATH)

    for d in (base_dir, to_long_unc(base_dir)):
        try:
            if not os.path.isdir(d): continue
            files = os.listdir(d)
            qgz = [f for f in files if f.lower().endswith(".qgz")]
            qgz = [f for f in qgz if normalize_name(f).startswith(normalize_name("Contexte éco -"))]
            return [os.path.join(d, f) for f in sorted(qgz)]
        except Exception:
            continue
    return []

# =========================
# Fonctions IGN (onglet 2) — identiques au script source
# =========================
def dms_to_dd(text: str) -> float:
    pat = r"(\d{1,3})[°d]\s*(\d{1,2})['m]\s*([\d\.]+)[\"s]?\s*([NSEW])"
    alt = r"(\d{1,3})\s+(\d{1,2})\s+([\d\.]+)\s*([NSEW])"
    m = re.search(pat, text, re.I) or re.search(alt, text, re.I)
    if not m:
        raise ValueError(f"Format DMS invalide : {text}")
    deg, mn, sc, hemi = m.groups()
    dd = float(deg) + float(mn)/60 + float(sc)/3600
    return -dd if hemi.upper() in ("S", "W") else dd

def add_hyperlink(paragraph, url: str, text: str, italic: bool = True):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    from docx.oxml import OxmlElement, ns
    fld_simple = OxmlElement('w:hyperlink')
    fld_simple.set(ns.qn('r:id'), r_id)

    run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    if italic:
        i = OxmlElement('w:i')
        r_pr.append(i)
    u = OxmlElement('w:u')
    u.set(ns.qn('w:val'), 'single')
    r_pr.append(u)
    run.append(r_pr)
    run_text = OxmlElement('w:t')
    run_text.text = text
    run.append(run_text)
    fld_simple.append(run)
    paragraph._p.append(fld_simple)

# =========================
# UI — Styles communs
# =========================
class StyleHelper:
    def __init__(self, master, prefs: dict):
        self.master = master
        self.prefs = prefs
        self.style = ttk.Style()
        try: self.style.theme_use("clam")
        except Exception: pass

    def apply(self, light: bool = True):
        if light:
            bg, fg, card_bg, accent, border, subfg = "#F6F7F9", "#111827", "#FFFFFF", "#2563EB", "#E5E7EB", "#6B7280"
        else:
            bg, fg, card_bg, accent, border, subfg = "#0F172A", "#E5E7EB", "#111827", "#3B82F6", "#1F2937", "#9CA3AF"

        self.master.configure(bg=bg)
        s = self.style
        s.configure(".", background=bg, foreground=fg, fieldbackground=card_bg, bordercolor=border)
        s.configure("Card.TFrame", background=card_bg, bordercolor=border, relief="solid", borderwidth=1)
        s.configure("Header.TFrame", background=card_bg, bordercolor=border, relief="flat")
        s.configure("TLabel", background=bg, foreground=fg)
        s.configure("Card.TLabel", background=card_bg, foreground=fg)
        s.configure("Subtle.TLabel", background=bg, foreground=subfg)
        s.configure("Tooltip.TLabel", background="#111827", foreground="#F9FAFB")
        s.configure("Accent.TButton", padding=10, background=accent, foreground="#FFFFFF")
        s.map("Accent.TButton", background=[("active", "#1D4ED8")], foreground=[("active", "#FFFFFF")])
        s.configure("Card.TCheckbutton", background=card_bg, foreground=fg)
        s.configure("Card.TRadiobutton", background=card_bg, foreground=fg)
        s.configure("Card.TEntry", fieldbackground=card_bg)
        s.configure("Status.TLabel", background=card_bg, foreground=subfg)
        s.configure("TProgressbar", troughcolor=border)

# =========================
# Onglet 1 — Export Cartes (mêmes fonctionnalités)
# =========================
class ExportCartesTab(ttk.Frame):
    def __init__(self, parent, style_helper: StyleHelper, prefs: dict):
        super().__init__(parent, padding=12)
        self.parent = parent
        self.prefs = prefs
        self.style_helper = style_helper

        self.font_title = tkfont.Font(family="Segoe UI", size=15, weight="bold")
        self.font_sub   = tkfont.Font(family="Segoe UI", size=10)
        self.font_mono  = tkfont.Font(family="Consolas", size=9)

        self.ze_shp_var   = tk.StringVar(value=self.prefs.get("ZE_SHP", ""))
        self.ae_shp_var   = tk.StringVar(value=self.prefs.get("AE_SHP", ""))
        self.cadrage_var  = tk.StringVar(value=self.prefs.get("CADRAGE_MODE", "BOTH"))
        self.overwrite_var= tk.BooleanVar(value=self.prefs.get("OVERWRITE", OVERWRITE_DEFAULT))
        self.dpi_var      = tk.IntVar(value=int(self.prefs.get("DPI", DPI_DEFAULT)))
        self.workers_var  = tk.IntVar(value=int(self.prefs.get("N_WORKERS", N_WORKERS_DEFAULT)))
        self.margin_var   = tk.DoubleVar(value=float(self.prefs.get("MARGIN_FAC", MARGIN_FAC_DEFAULT)))

        self.project_vars: dict[str, tk.IntVar] = {}
        self.all_projects: List[str] = []
        self.filtered_projects: List[str] = []
        self.total_expected = 0
        self.progress_done  = 0

        self._build_ui()
        self._populate_projects()
        self._update_counts()

    def _build_ui(self):
        header = ttk.Frame(self, style="Header.TFrame", padding=(14, 12))
        header.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(header, text="Export cartes — QGIS → PNG", style="Card.TLabel", font=self.font_title)\
            .grid(row=0, column=0, sticky="w")
        ttk.Label(header, text="Sélection shapefiles, choix du cadrage, export multi-projets.", style="Subtle.TLabel", font=self.font_sub)\
            .grid(row=1, column=0, sticky="w", pady=(4,0))
        header.columnconfigure(0, weight=1)

        grid = ttk.Frame(self); grid.pack(fill=tk.BOTH, expand=True)
        left = ttk.Frame(grid);  left.grid(row=0, column=0, sticky="nsew", padx=(0, 8))
        right = ttk.Frame(grid); right.grid(row=0, column=1, sticky="nsew", padx=(8, 0))
        grid.columnconfigure(0, weight=1); grid.columnconfigure(1, weight=1); grid.rowconfigure(0, weight=1)

        # Shapefiles
        shp = ttk.Frame(left, style="Card.TFrame", padding=12); shp.pack(fill=tk.X)
        ttk.Label(shp, text="1. Couches Shapefile", style="Card.TLabel").grid(row=0, column=0, columnspan=4, sticky="w")
        self._file_row(shp, 1, "📁 Zone d'étude…", self.ze_shp_var, lambda: self._select_shapefile('ZE'))
        self._file_row(shp, 2, "📁 Aire d'étude élargie…", self.ae_shp_var, lambda: self._select_shapefile('AE'))
        shp.columnconfigure(1, weight=1)

        # Options
        opt = ttk.Frame(left, style="Card.TFrame", padding=12); opt.pack(fill=tk.X, pady=(10,0))
        ttk.Label(opt, text="2. Cadrage et options", style="Card.TLabel").grid(row=0, column=0, columnspan=6, sticky="w")
        ttk.Radiobutton(opt, text="AE + ZE", variable=self.cadrage_var, value="BOTH", style="Card.TRadiobutton").grid(row=1, column=0, sticky="w", pady=(6,2))
        ttk.Radiobutton(opt, text="ZE uniquement", variable=self.cadrage_var, value="ZE", style="Card.TRadiobutton").grid(row=1, column=1, sticky="w", padx=(12,0))
        ttk.Radiobutton(opt, text="AE uniquement", variable=self.cadrage_var, value="AE", style="Card.TRadiobutton").grid(row=1, column=2, sticky="w", padx=(12,0))
        ttk.Checkbutton(opt, text="Écraser si le PNG existe", variable=self.overwrite_var, style="Card.TCheckbutton").grid(row=1, column=3, sticky="w", padx=(24, 0))

        ttk.Label(opt, text="DPI", style="Card.TLabel").grid(row=2, column=0, sticky="w", pady=(8,0))
        ttk.Spinbox(opt, from_=72, to=1200, textvariable=self.dpi_var, width=6, justify="right").grid(row=2, column=1, sticky="w", pady=(8,0))
        ttk.Label(opt, text="Workers", style="Card.TLabel").grid(row=2, column=2, sticky="w", padx=(12,0), pady=(8,0))
        ttk.Spinbox(opt, from_=1, to=max(1, (os.cpu_count() or 2)), textvariable=self.workers_var, width=6, justify="right").grid(row=2, column=3, sticky="w", pady=(8,0))
        ttk.Label(opt, text="Marge", style="Card.TLabel").grid(row=2, column=4, sticky="w", padx=(12,0), pady=(8,0))
        ttk.Spinbox(opt, from_=1.00, to=2.00, increment=0.05, textvariable=self.margin_var, width=6, justify="right").grid(row=2, column=5, sticky="w", pady=(8,0))

        # Actions
        act = ttk.Frame(left, style="Card.TFrame", padding=12); act.pack(fill=tk.X, pady=(10,0))
        self.export_button = ttk.Button(act, text="▶ Lancer l’export", style="Accent.TButton", command=self.start_export_thread)
        self.export_button.grid(row=0, column=0, sticky="w")
        obtn = ttk.Button(act, text="📂 Ouvrir le dossier de sortie", command=self._open_out_dir)
        obtn.grid(row=0, column=1, padx=(10,0)); ToolTip(obtn, OUT_IMG)
        tbtn = ttk.Button(act, text="🧪 Tester QGIS", command=self._test_qgis_threaded)
        tbtn.grid(row=0, column=2, padx=(10,0)); ToolTip(tbtn, "Vérifier l’import QGIS/Qt")

        # Projets
        proj = ttk.Frame(right, style="Card.TFrame", padding=12); proj.pack(fill=tk.BOTH, expand=True)
        ttk.Label(proj, text="3. Projets QGIS", style="Card.TLabel").grid(row=0, column=0, columnspan=4, sticky="w")
        ttk.Label(proj, text="Filtrer", style="Card.TLabel").grid(row=1, column=0, sticky="w", pady=(6,6))
        self.filter_var = tk.StringVar()
        fe = ttk.Entry(proj, textvariable=self.filter_var, width=32); fe.grid(row=1, column=1, sticky="w", pady=(6,6))
        fe.bind("<KeyRelease>", lambda _e: self._apply_filter())
        ttk.Button(proj, text="Tout", width=6, command=lambda: self._select_all(True)).grid(row=1, column=2, padx=(8,0))
        ttk.Button(proj, text="Aucun", width=6, command=lambda: self._select_all(False)).grid(row=1, column=3, padx=(6,0))

        canvas = tk.Canvas(proj, highlightthickness=0, borderwidth=0)
        scrollbar = ttk.Scrollbar(proj, orient="vertical", command=canvas.yview)
        self.scrollable_frame = ttk.Frame(canvas)
        self.scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.grid(row=2, column=0, columnspan=4, sticky="nsew", pady=(6, 6))
        scrollbar.grid(row=2, column=4, sticky="ns", padx=(6,0))
        proj.rowconfigure(2, weight=1); proj.columnconfigure(1, weight=1)

        # Bas
        bottom = ttk.Frame(self, style="Card.TFrame", padding=12); bottom.pack(fill=tk.BOTH, expand=True, pady=(10,0))
        self.status_label = ttk.Label(bottom, text="Prêt.", style="Status.TLabel"); self.status_label.grid(row=0, column=0, sticky="w")
        self.progress = ttk.Progressbar(bottom, orient="horizontal", mode="determinate", length=220)
        self.progress.grid(row=0, column=1, sticky="e"); bottom.columnconfigure(0, weight=1)

        log_frame = ttk.Frame(bottom, style="Card.TFrame"); log_frame.grid(row=1, column=0, columnspan=2, sticky="nsew", pady=(8,0))
        bottom.rowconfigure(1, weight=1)
        self.log_text = tk.Text(log_frame, height=10, wrap=tk.WORD, state='disabled',
                                bg=self.style_helper.style.lookup("Card.TFrame", "background"),
                                fg=self.style_helper.style.lookup("TLabel", "foreground"))
        self.log_text.configure(font=self.font_mono, relief="flat")
        log_scroll = ttk.Scrollbar(log_frame, orient="vertical", command=self.log_text.yview)
        self.log_text['yscrollcommand'] = log_scroll.set
        log_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sys.stdout = TextRedirector(self.log_text)

    def _file_row(self, parent, row: int, label: str, var: tk.StringVar, cmd):
        btn = ttk.Button(parent, text=label, command=cmd)
        btn.grid(row=row, column=0, sticky="w", pady=(8 if row == 1 else 4, 2))
        ent = ttk.Entry(parent, textvariable=var, width=10); ent.grid(row=row, column=1, sticky="ew", padx=8)
        ent.configure(state="readonly")
        copy_btn = ttk.Button(parent, text="Copier", width=6, command=lambda: self._copy_to_clipboard(var.get()))
        copy_btn.grid(row=row, column=2, sticky="e")
        clear_btn = ttk.Button(parent, text="✖", width=3, command=lambda: var.set(""))
        clear_btn.grid(row=row, column=3, sticky="e")
        ToolTip(copy_btn, "Copier le chemin"); ToolTip(clear_btn, "Effacer")
        parent.columnconfigure(1, weight=1)

    def _copy_to_clipboard(self, text: str):
        try: self.winfo_toplevel().clipboard_clear(); self.winfo_toplevel().clipboard_append(text)
        except Exception: pass

    def _select_shapefile(self, shp_type):
        label_text = "Zone d'étude" if shp_type == 'ZE' else "Aire d'étude élargie"
        title = f"Sélectionner le shapefile pour '{label_text}'"
        base_dir = DEFAULT_SHAPE_DIR if os.path.isdir(DEFAULT_SHAPE_DIR) else os.path.expanduser("~")
        path = filedialog.askopenfilename(title=title, initialdir=base_dir, filetypes=[("Shapefile ESRI", "*.shp")])
        if path:
            if shp_type == 'ZE': self.ze_shp_var.set(path)
            else: self.ae_shp_var.set(path)

    def _populate_projects(self):
        for w in list(self.scrollable_frame.children.values()): w.destroy()
        self.project_vars.clear()
        self.all_projects = discover_projects()
        self.filtered_projects = list(self.all_projects)
        if not self.all_projects:
            ttk.Label(self.scrollable_frame, text="Aucun projet trouvé ou dossier inaccessible.", foreground="red").pack(anchor="w")
            return
        for proj_path in self.filtered_projects:
            var = tk.IntVar(value=1); self.project_vars[proj_path] = var
            ttk.Checkbutton(self.scrollable_frame, text=os.path.basename(proj_path), variable=var, style="Card.TCheckbutton")\
                .pack(anchor='w', padx=4, pady=1)

    def _apply_filter(self):
        term = normalize_name(self.filter_var.get())
        for w in list(self.scrollable_frame.children.values()): w.destroy()
        self.filtered_projects = [p for p in self.all_projects if term in normalize_name(os.path.basename(p))]
        if not self.filtered_projects:
            ttk.Label(self.scrollable_frame, text="Aucun projet ne correspond au filtre.", foreground="red").pack(anchor="w")
            self.project_vars = {}; self._update_counts(); return
        for proj_path in self.filtered_projects:
            current = self.project_vars.get(proj_path, tk.IntVar(value=1))
            self.project_vars[proj_path] = current
            ttk.Checkbutton(self.scrollable_frame, text=os.path.basename(proj_path), variable=current, style="Card.TCheckbutton")\
                .pack(anchor='w', padx=4, pady=1)
        self._update_counts()

    def _select_all(self, state: bool):
        for var in self.project_vars.values(): var.set(1 if state else 0)
        self._update_counts()

    def _selected_projects(self) -> List[str]:
        return [p for p, v in self.project_vars.items() if v.get() == 1 and p in self.filtered_projects]

    def _update_counts(self):
        selected = len(self._selected_projects()); total = len(self.filtered_projects)
        self.status_label.config(text=f"Projets sélectionnés : {selected} / {total}")

    def _open_out_dir(self):
        try:
            os.makedirs(OUT_IMG, exist_ok=True); os.startfile(OUT_IMG)
        except Exception as e:
            messagebox.showerror("Erreur", f"Impossible d’ouvrir le dossier de sortie : {e}")

    def _test_qgis_threaded(self):
        t = threading.Thread(target=self._test_qgis); t.daemon = True; t.start()

    def _test_qgis(self):
        try:
            log_with_time("Test QGIS : import…")
            qt_base = None
            for name in ("Qt6", "Qt5"):
                base = os.path.join(QGIS_ROOT, "apps", name)
                if os.path.isdir(base): qt_base = base; break
            if not qt_base: raise RuntimeError("Répertoire Qt introuvable")
            sys.path.insert(0, os.path.join(QGIS_APP, "python"))
            sys.path.insert(0, os.path.join(QGIS_ROOT, "apps", PY_VER, "Lib", "site-packages"))
            from qgis.core import QgsApplication  # noqa: F401
            log_with_time("Test QGIS : OK")
            messagebox.showinfo("QGIS", "Import QGIS OK.")
        except Exception as e:
            log_with_time(f"Échec import QGIS : {e}")
            messagebox.showerror("QGIS", f"Échec import QGIS : {e}")

    def start_export_thread(self):
        if not self.ze_shp_var.get() or not self.ae_shp_var.get():
            messagebox.showerror("Erreur", "Sélectionnez les deux shapefiles."); return
        if not os.path.isfile(self.ze_shp_var.get()) or not os.path.isfile(self.ae_shp_var.get()):
            messagebox.showerror("Erreur", "Un shapefile est introuvable."); return
        projets = self._selected_projects()
        if not projets:
            messagebox.showerror("Erreur", "Sélectionnez au moins un projet."); return

        self._update_counts()
        self.export_button.config(state="disabled")
        self.progress_done = 0; self.progress["value"] = 0

        mode = self.cadrage_var.get(); per_project = 2 if mode == "BOTH" else 1
        self.total_expected = per_project * len(projets)
        self.progress["maximum"] = max(1, self.total_expected)

        self.prefs.update({
            "ZE_SHP": self.ze_shp_var.get(),
            "AE_SHP": self.ae_shp_var.get(),
            "CADRAGE_MODE": self.cadrage_var.get(),
            "OVERWRITE": bool(self.overwrite_var.get()),
            "DPI": int(self.dpi_var.get()),
            "N_WORKERS": int(self.workers_var.get()),
            "MARGIN_FAC": float(self.margin_var.get()),
        }); save_prefs(self.prefs)

        t = threading.Thread(target=self._run_export_logic, args=(projets,))
        t.daemon = True; t.start()

    def _run_export_logic(self, projets: List[str]):
        try:
            start = datetime.datetime.now()
            os.makedirs(OUT_IMG, exist_ok=True)

            log_with_time(f"{len(projets)} projets (attendu = calcul en cours)")
            log_with_time(f"Workers={self.workers_var.get()}, DPI={self.dpi_var.get()}, marge={self.margin_var.get():.2f}, overwrite={self.overwrite_var.get()}")

            chunks = chunk_even(projets, self.workers_var.get())
            cfg = {
                "QGIS_ROOT": QGIS_ROOT, "QGIS_APP": QGIS_APP, "PY_VER": PY_VER,
                "OUT_IMG": OUT_IMG, "DPI": int(self.dpi_var.get()),
                "MARGIN_FAC": float(self.margin_var.get()),
                "LAYER_AE_NAME": LAYER_AE_NAME, "LAYER_ZE_NAME": LAYER_ZE_NAME,
                "AE_SHP": self.ae_shp_var.get(), "ZE_SHP": self.ze_shp_var.get(),
                "CADRAGE_MODE": self.cadrage_var.get(), "OVERWRITE": bool(self.overwrite_var.get()),
            }

            ok_total = 0; ko_total = 0
            def ui_update_progress(done_inc):
                self.progress_done += done_inc
                self.progress["value"] = min(self.progress_done, self.total_expected)
                self.status_label.config(text=f"Progression : {self.progress_done}/{self.total_expected}")

            with ProcessPoolExecutor(max_workers=int(self.workers_var.get())) as ex:
                futures = [ex.submit(worker_run, (chunk, cfg)) for chunk in chunks if chunk]
                for fut in as_completed(futures):
                    try:
                        ok, ko = fut.result()
                        ok_total += ok; ko_total += ko
                        self.after(0, ui_update_progress, ok + ko)
                        log_with_time(f"Lot terminé: {ok} OK, {ko} KO")
                    except Exception as e:
                        log_with_time(f"Erreur worker: {e}")

            elapsed = datetime.datetime.now() - start
            log_with_time(f"FIN — OK={ok_total} | KO={ko_total} | Attendu={self.total_expected} | Durée={elapsed}")
            self.after(0, lambda: self.status_label.config(text=f"Terminé — OK={ok_total} / KO={ko_total}"))
        except Exception as e:
            log_with_time(f"Erreur critique: {e}")
            self.after(0, lambda: messagebox.showerror("Erreur", str(e)))
        finally:
            self.after(0, lambda: self.export_button.config(state="normal"))

# =========================
# Onglet 2 — Remonter le temps (UI + logique)
# =========================
class RemonterLeTempsTab(ttk.Frame):
    def __init__(self, parent, style_helper: StyleHelper, prefs: dict):
        super().__init__(parent, padding=12)
        self.style_helper = style_helper
        self.prefs = prefs

        self.font_title = tkfont.Font(family="Segoe UI", size=15, weight="bold")
        self.font_sub   = tkfont.Font(family="Segoe UI", size=10)
        self.font_mono  = tkfont.Font(family="Consolas", size=9)

        self.coord_var   = tk.StringVar(value=self.prefs.get("RLT_COORD", ""))   # ex: 45°09'30" N 5°43'12" E
        self.commune_var = tk.StringVar(value=self.prefs.get("RLT_COMMUNE", ""))
        self.wait_var    = tk.DoubleVar(value=float(self.prefs.get("RLT_WAIT", WAIT_TILES_DEFAULT)))
        self.out_dir_var = tk.StringVar(value=self.prefs.get("RLT_OUT", OUTPUT_DIR_RLT))
        self.headless_var= tk.BooleanVar(value=bool(self.prefs.get("RLT_HEADLESS", False)))

        self._build_ui()

    def _build_ui(self):
        header = ttk.Frame(self, style="Header.TFrame", padding=(14, 12))
        header.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(header, text="IGN « Remonter le temps » — Capture + Word", style="Card.TLabel", font=self.font_title)\
            .grid(row=0, column=0, sticky="w")
        ttk.Label(header, text="Entrer coordonnées DMS et commune. Générer 2×2 + commentaire.", style="Subtle.TLabel", font=self.font_sub)\
            .grid(row=1, column=0, sticky="w", pady=(4,0))
        header.columnconfigure(0, weight=1)

        # Carte paramètres
        card = ttk.Frame(self, style="Card.TFrame", padding=12)
        card.pack(fill=tk.X)
        r = 0
        ttk.Label(card, text="Coordonnées DMS (lat puis lon)", style="Card.TLabel").grid(row=r, column=0, sticky="w")
        ttk.Entry(card, textvariable=self.coord_var).grid(row=r, column=1, sticky="ew", padx=8)
        ToolTip(card, "Exemple : 45°09'30\" N 5°43'12\" E")
        r += 1

        ttk.Label(card, text="Commune", style="Card.TLabel").grid(row=r, column=0, sticky="w")
        ttk.Entry(card, textvariable=self.commune_var).grid(row=r, column=1, sticky="ew", padx=8)
        r += 1

        ttk.Label(card, text="Dossier de sortie", style="Card.TLabel").grid(row=r, column=0, sticky="w")
        out_row = ttk.Frame(card, style="Card.TFrame")
        out_row.grid(row=r, column=1, sticky="ew", padx=0)
        out_row.columnconfigure(0, weight=1)
        ttk.Entry(out_row, textvariable=self.out_dir_var).grid(row=0, column=0, sticky="ew")
        ttk.Button(out_row, text="Parcourir…", command=self._pick_out_dir).grid(row=0, column=1, padx=(6,0))
        r += 1

        ttk.Label(card, text="Attente tuiles (s)", style="Card.TLabel").grid(row=r, column=0, sticky="w")
        ttk.Spinbox(card, from_=0.5, to=5.0, increment=0.1, textvariable=self.wait_var, width=6, justify="right").grid(row=r, column=1, sticky="w", padx=8)
        r += 1

        ttk.Checkbutton(card, text="Mode headless (Chrome sans fenêtre)", variable=self.headless_var, style="Card.TCheckbutton")\
            .grid(row=r, column=1, sticky="w", padx=8, pady=(4,0))
        r += 1

        card.columnconfigure(1, weight=1)

        # Actions
        act = ttk.Frame(self, style="Card.TFrame", padding=12)
        act.pack(fill=tk.X, pady=(10,0))
        self.run_btn = ttk.Button(act, text="▶ Capturer + Générer le Word", style="Accent.TButton", command=self._start_thread)
        self.run_btn.grid(row=0, column=0, sticky="w")
        obtn = ttk.Button(act, text="📂 Ouvrir le dossier de sortie", command=self._open_out_dir)
        obtn.grid(row=0, column=1, padx=(10,0)); ToolTip(obtn, "Ouvrir le dossier cible")

        # Logs
        bottom = ttk.Frame(self, style="Card.TFrame", padding=12)
        bottom.pack(fill=tk.BOTH, expand=True, pady=(10,0))
        self.status_label = ttk.Label(bottom, text="Prêt.", style="Status.TLabel")
        self.status_label.grid(row=0, column=0, sticky="w")
        bottom.columnconfigure(0, weight=1)

        log_frame = ttk.Frame(bottom, style="Card.TFrame")
        log_frame.grid(row=1, column=0, sticky="nsew", pady=(8,0))
        bottom.rowconfigure(1, weight=1)

        self.log_text = tk.Text(log_frame, height=12, wrap=tk.WORD, state='disabled',
                                bg=self.style_helper.style.lookup("Card.TFrame", "background"),
                                fg=self.style_helper.style.lookup("TLabel", "foreground"))
        self.log_text.configure(font=self.font_mono, relief="flat")
        log_scroll = ttk.Scrollbar(log_frame, orient="vertical", command=self.log_text.yview)
        self.log_text['yscrollcommand'] = log_scroll.set
        log_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.stdout_redirect = TextRedirector(self.log_text)  # dédié à l’onglet
        # Utilisation: print(..., file=self.stdout_redirect)

    # --- Actions onglet 2 ---
    def _pick_out_dir(self):
        base = self.out_dir_var.get() or OUTPUT_DIR_RLT
        d = filedialog.askdirectory(title="Choisir le dossier de sortie", initialdir=base if os.path.isdir(base) else os.path.expanduser("~"))
        if d:
            self.out_dir_var.set(d)

    def _open_out_dir(self):
        try:
            os.makedirs(self.out_dir_var.get() or OUTPUT_DIR_RLT, exist_ok=True)
            os.startfile(self.out_dir_var.get() or OUTPUT_DIR_RLT)
        except Exception as e:
            messagebox.showerror("Erreur", f"Impossible d’ouvrir le dossier : {e}")

    def _start_thread(self):
        if not self.coord_var.get().strip():
            messagebox.showerror("Erreur", "Renseigner les coordonnées en DMS."); return
        if not self.commune_var.get().strip():
            messagebox.showerror("Erreur", "Renseigner le nom de la commune."); return
        self.run_btn.config(state="disabled")
        t = threading.Thread(target=self._run_process)
        t.daemon = True
        t.start()

    def _run_process(self):
        try:
            # Sauvegarde préférences
            self.prefs.update({
                "RLT_COORD": self.coord_var.get().strip(),
                "RLT_COMMUNE": self.commune_var.get().strip(),
                "RLT_WAIT": float(self.wait_var.get()),
                "RLT_OUT": self.out_dir_var.get().strip(),
                "RLT_HEADLESS": bool(self.headless_var.get()),
            }); save_prefs(self.prefs)

            print(f"[IGN] Parsing coordonnées…", file=self.stdout_redirect)
            # La chaîne doit contenir LAT puis LON en DMS, séparés par espace/virgule/tab
            parts = re.split(r"\s{2,}|,|\t", self.coord_var.get().strip())
            if len(parts) < 2:
                # fallback : tente un split sur premier espace
                parts = re.split(r"\s+", self.coord_var.get().strip(), maxsplit=1)
            if len(parts) < 2:
                raise ValueError("Coordonnées DMS attendues au format « LAT  LON »")

            lat_dd = dms_to_dd(parts[0])
            lon_dd = dms_to_dd(parts[1])
            wait_s = float(self.wait_var.get())
            out_dir = self.out_dir_var.get().strip() or OUTPUT_DIR_RLT
            os.makedirs(out_dir, exist_ok=True)
            commune = self.commune_var.get().strip()
            comment_txt = COMMENT_TEMPLATE.format(commune=commune)

            drv_opts = webdriver.ChromeOptions()
            drv_opts.add_argument("--log-level=3")
            drv_opts.add_experimental_option('excludeSwitches', ['enable-logging'])
            drv_opts.add_argument("--disable-extensions")
            if self.headless_var.get():
                drv_opts.add_argument("--headless=new")

            print(f"[IGN] Lancement Chrome…", file=self.stdout_redirect)
            driver = webdriver.Chrome(options=drv_opts)
            try:
                driver.maximize_window()
            except Exception:
                pass

            images = []
            viewport = (By.CSS_SELECTOR, "div.ol-viewport")

            for title, layer_val in LAYERS:
                url = URL.format(lon=f"{lon_dd:.6f}", lat=f"{lat_dd:.6f}", layer=layer_val)
                print(f"[IGN] {title} → {url}", file=self.stdout_redirect)
                driver.get(url)
                WebDriverWait(driver, 20).until(EC.visibility_of_element_located(viewport))
                time.sleep(wait_s)
                tgt = driver.find_element(*viewport)
                img_path = os.path.join(out_dir, f"{title}.png")
                if tgt.screenshot(img_path):
                    img = Image.open(img_path)
                    w, h = img.size
                    left, right = int(w * 0.05), int(w * 0.95)
                    img.crop((left, 0, right, h)).save(img_path)
                    images.append((title, img_path))
                    print(f"[IGN] Capture OK : {img_path}", file=self.stdout_redirect)
                else:
                    print(f"[IGN] Capture échouée : {title}", file=self.stdout_redirect)

            driver.quit()

            if not images:
                print("[IGN] Aucune image → pas de doc.", file=self.stdout_redirect)
                messagebox.showwarning("IGN", "Aucune image capturée.")
                return

            print("[IGN] Génération du Word…", file=self.stdout_redirect)
            doc = Document()
            style_normal = doc.styles['Normal']
            style_normal.font.name = 'Calibri'
            style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            style_tbl = doc.styles['Table Grid']
            style_tbl.font.name = 'Calibri'
            style_tbl._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

            sec = doc.sections[0]
            sec.orientation = WD_ORIENT.LANDSCAPE
            sec.page_width, sec.page_height = sec.page_height, sec.page_width
            for m in (sec.left_margin, sec.right_margin, sec.top_margin, sec.bottom_margin):
                m = Cm(1.5)

            cap_par = doc.add_paragraph()
            cap_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_hyperlink(cap_par, "https://remonterletemps.ign.fr/",
                          f"Comparaison temporelle — {commune} (source : IGN – RemonterLeTemps)")

            table = doc.add_table(rows=2, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            table.autofit = False

            for idx, (title, path) in enumerate(images):
                r, c = divmod(idx, 2)
                cell = table.cell(r, c)
                p_t = cell.paragraphs[0]
                run_t = p_t.add_run(title); run_t.bold = True
                p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.add_paragraph()
                p_img = cell.add_paragraph()
                if os.path.exists(path):
                    p_img.add_run().add_picture(path, width=IMG_WIDTH)
                else:
                    p_img.add_run(f"[image manquante : {title}]")
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()
            p_comm = doc.add_paragraph(comment_txt)
            p_comm.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            doc_path = os.path.join(out_dir, WORD_FILENAME)
            doc.save(doc_path)
            print(f"[IGN] Document généré : {doc_path}", file=self.stdout_redirect)
            self._set_status(f"Terminé — {doc_path}")
        except Exception as e:
            print(f"[IGN] Erreur : {e}", file=self.stdout_redirect)
            messagebox.showerror("IGN", str(e))
        finally:
            self.after(0, lambda: self.run_btn.config(state="normal"))

    def _set_status(self, txt: str):
        self.after(0, lambda: self.status_label.config(text=txt))

# =========================
# Onglet 3 — Identification Plantes (Pl@ntNet)
# =========================

# --- CONFIGURATION ---
API_KEY = "2b10vfT6MvFC2lcAzqG1ZMKO"  # Votre clé API Pl@ntNet
PROJECT = "all"  # Vous pouvez spécifier une flore particulière comme "weurope", "canada", etc.
API_URL = f"https://my-api.plantnet.org/v2/identify/{PROJECT}?api-key={API_KEY}"

FOLDER = ""  # Sélectionné par l'utilisateur via l'onglet
# --- FIN CONFIGURATION ---

# Fonction pour réduire la taille et compresser l'image
def resize_image(image_path, max_size=(800, 800), quality=70):
    """
    Redimensionne et compresse une image.

    :param image_path: Chemin de l'image à traiter.
    :param max_size: Tuple indiquant la taille maximale (largeur, hauteur).
    :param quality: Qualité de compression (1-100).
    :return: BytesIO de l'image traitée ou None en cas d'erreur.
    """
    try:
        with Image.open(image_path) as img:
            img.thumbnail(max_size)
            buffer = BytesIO()
            img.save(buffer, format='JPEG', quality=quality)  # Convertir en JPEG pour une meilleure compression
            buffer.seek(0)
            return buffer
    except Exception as e:
        print(f"Erreur lors du redimensionnement de l'image : {e}")
        return None

# Fonction pour identifier la plante via l'API Pl@ntNet
def identify_plant(image_path, organ):
    """
    Envoie une image à l'API Pl@ntNet pour identification.

    :param image_path: Chemin de l'image à envoyer.
    :param organ: Type d'organe de la plante (par exemple, 'flower').
    :return: Nom scientifique de la plante identifiée ou None.
    """
    print(f"Envoi de l'image à l'API : {image_path}")
    try:
        resized_image = resize_image(image_path)
        if not resized_image:
            print(f"Échec du redimensionnement de l'image : {image_path}")
            return None

        files = {
            'images': (os.path.basename(image_path), resized_image, 'image/jpeg')
        }
        data = {
            'organs': organ
        }

        # Envoyer la requête
        response = requests.post(API_URL, files=files, data=data)

        print(f"Réponse de l'API : {response.status_code}")
        if response.status_code == 200:
            json_result = response.json()
            try:
                species = json_result['results'][0]['species']['scientificNameWithoutAuthor']
                print(f"Plante identifiée : {species}")
                return species
            except (KeyError, IndexError):
                print(f"Aucun résultat trouvé pour l'image : {image_path}")
                return None
        else:
            print(f"Erreur API : {response.status_code} - {response.text}")
            return None
    except Exception as e:
        print(f"Exception lors de l'identification de la plante : {e}")
        return None

# Fonction pour copier et renommer le fichier
def copy_and_rename_file(file_path, new_name, count):
    """
    Copie et renomme un fichier dans le dossier de destination.

    :param file_path: Chemin du fichier original.
    :param new_name: Nom scientifique de la plante.
    :param count: Compteur pour différencier les fichiers portant le même nom.
    """
    ext = os.path.splitext(file_path)[1]
    if count == 1:
        # Premier fichier avec ce nom, pas de numéro
        new_file_name = f"{new_name} @plantnet{ext}"
    else:
        # Ajouter un numéro entre parenthèses après '@plantnet'
        new_file_name = f"{new_name} @plantnet({count}){ext}"
    new_path = os.path.join(FOLDER, new_file_name)
    try:
        shutil.copy(file_path, new_path)
        print(f"Fichier copié et renommé : {file_path} -> {new_path}")
    except Exception as e:
        print(f"Erreur lors de la copie du fichier : {e}")


class PlantNetTab(ttk.Frame):
    def __init__(self, parent, style_helper: StyleHelper, prefs: dict):
        super().__init__(parent, padding=12)
        self.style_helper = style_helper
        self.prefs = prefs

        self.font_title = tkfont.Font(family="Segoe UI", size=15, weight="bold")
        self.font_sub   = tkfont.Font(family="Segoe UI", size=10)
        self.font_mono  = tkfont.Font(family="Consolas", size=9)

        self.folder_var = tk.StringVar()

        self._build_ui()

    def _build_ui(self):
        header = ttk.Frame(self, style="Header.TFrame", padding=(14, 12))
        header.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(header, text="Pl@ntNet — Identification", style="Card.TLabel", font=self.font_title)\
            .grid(row=0, column=0, sticky="w")
        ttk.Label(header, text="Sélectionner un dossier d'images puis lancer l'analyse.", style="Subtle.TLabel", font=self.font_sub)\
            .grid(row=1, column=0, sticky="w", pady=(4,0))
        header.columnconfigure(0, weight=1)

        card = ttk.Frame(self, style="Card.TFrame", padding=12)
        card.pack(fill=tk.X)
        ttk.Label(card, text="Dossier d'images", style="Card.TLabel").grid(row=0, column=0, sticky="w")
        ttk.Entry(card, textvariable=self.folder_var).grid(row=0, column=1, sticky="ew", padx=(8,0))
        ttk.Button(card, text="Parcourir…", command=self._select_folder).grid(row=0, column=2, padx=(8,0))
        card.columnconfigure(1, weight=1)

        act = ttk.Frame(self, style="Card.TFrame", padding=12)
        act.pack(fill=tk.X, pady=(10,0))
        ttk.Button(act, text="▶ Lancer l’identification", style="Accent.TButton", command=self._start_thread)\
            .grid(row=0, column=0, sticky="w")

        bottom = ttk.Frame(self, style="Card.TFrame", padding=12)
        bottom.pack(fill=tk.BOTH, expand=True, pady=(10,0))
        log_frame = ttk.Frame(bottom, style="Card.TFrame")
        log_frame.pack(fill=tk.BOTH, expand=True)
        self.log_text = tk.Text(log_frame, height=12, wrap=tk.WORD, state='disabled',
                                bg=self.style_helper.style.lookup("Card.TFrame", "background"),
                                fg=self.style_helper.style.lookup("TLabel", "foreground"))
        self.log_text.configure(font=self.font_mono, relief="flat")
        log_scroll = ttk.Scrollbar(log_frame, orient="vertical", command=self.log_text.yview)
        self.log_text['yscrollcommand'] = log_scroll.set
        log_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.stdout_redirect = TextRedirector(self.log_text)

    def _select_folder(self):
        d = filedialog.askdirectory(title="Choisir le dossier d'images", initialdir=os.path.expanduser("~"))
        if d:
            self.folder_var.set(d)

    def _start_thread(self):
        if not self.folder_var.get():
            messagebox.showerror("Erreur", "Sélectionnez un dossier d'images.")
            return
        global FOLDER
        FOLDER = self.folder_var.get()
        t = threading.Thread(target=self._run_logic)
        t.daemon = True
        t.start()

    def _run_logic(self):
        old_stdout = sys.stdout
        sys.stdout = self.stdout_redirect
        try:
            if not API_KEY:
                print("Veuillez configurer votre clé API avant de lancer le script.")
            elif not os.path.exists(FOLDER):
                print(f"Le dossier à traiter n'existe pas : {FOLDER}")
            else:
                image_extensions = ['.jpg', '.jpeg', '.png', '.heic', '.heif']
                image_files = []
                for root_dir, dirs, files in os.walk(FOLDER):
                    for f in files:
                        if os.path.splitext(f)[1].lower() in image_extensions:
                            if '@plantnet' not in f:
                                image_files.append(os.path.join(root_dir, f))

                if not image_files:
                    print("Aucune image à traiter dans le dossier.")
                else:
                    plant_name_counts = {}
                    for image_path in image_files:
                        organ = 'flower'
                        plant_name = identify_plant(image_path, organ)
                        if plant_name:
                            count = plant_name_counts.get(plant_name, 0) + 1
                            plant_name_counts[plant_name] = count
                            copy_and_rename_file(image_path, plant_name, count)
                        else:
                            print(f"Aucune identification possible pour l'image : {image_path}")
        finally:
            sys.stdout = old_stdout


# =========================
# App principale avec Notebook
# =========================
class MainApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Contexte éco — Outils")
        self.root.geometry("1060x760"); self.root.minsize(900, 640)

        self.prefs = load_prefs()
        self.style_helper = StyleHelper(root, self.prefs)
        self.theme_var = tk.StringVar(value=self.prefs.get("theme", "light"))
        self.style_helper.apply(light=(self.theme_var.get() == "light"))

        # Header global + bouton thème
        top = ttk.Frame(root, style="Header.TFrame", padding=(12, 8))
        top.pack(fill=tk.X)
        ttk.Label(top, text="Contexte éco — Suite d’outils", style="Card.TLabel",
                  font=tkfont.Font(family="Segoe UI", size=16, weight="bold")).pack(side=tk.LEFT)
        btn_theme = ttk.Button(top, text="Thème clair/sombre", command=self._toggle_theme)
        btn_theme.pack(side=tk.RIGHT)

        # Notebook
        nb = ttk.Notebook(root)
        nb.pack(fill=tk.BOTH, expand=True, padx=12, pady=10)

        self.tab_export = ExportCartesTab(nb, self.style_helper, self.prefs)
        self.tab_rlt    = RemonterLeTempsTab(nb, self.style_helper, self.prefs)
        self.tab_plnt   = PlantNetTab(nb, self.style_helper, self.prefs)

        nb.add(self.tab_export, text="Export Cartes")
        nb.add(self.tab_rlt, text="Remonter le temps")
        nb.add(self.tab_plnt, text="Pl@ntNet")

        # Raccourcis utiles
        root.bind("<Control-1>", lambda _e: nb.select(0))
        root.bind("<Control-2>", lambda _e: nb.select(1))
        root.bind("<Control-3>", lambda _e: nb.select(2))

        # Sauvegarde prefs à la fermeture
        root.protocol("WM_DELETE_WINDOW", self._on_close)

    def _toggle_theme(self):
        new_theme = "dark" if self.theme_var.get() == "light" else "light"
        self.theme_var.set(new_theme)
        self.prefs["theme"] = new_theme
        save_prefs(self.prefs)
        self.style_helper.apply(light=(new_theme == "light"))

    def _on_close(self):
        try:
            save_prefs(self.prefs)
        finally:
            self.root.destroy()

# =========================
# Main
# =========================
if __name__ == "__main__":
    root = tk.Tk()
    app = MainApp(root)
    root.mainloop()
